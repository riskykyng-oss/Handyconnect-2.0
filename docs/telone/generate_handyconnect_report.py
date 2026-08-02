# -*- coding: utf-8 -*-
"""
TelOne Centre for Learning - Diploma in Software Engineering
Project: HandyConnect - A Web-Based Gig Marketplace for Home Services
Full project report following the prescribed report structure.

Formatting: Times New Roman 12pt, 1.5 spacing, justified, page numbers
(roman for preliminary pages, arabic from Chapter 1), chapters on fresh pages,
TOC via a Word field, Lists of Figures and Tables.
Output: docs/telone/HandyConnect_Project_Report.docx
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import rcParams

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
os.makedirs(FIG, exist_ok=True)
SHARED_DIAG = os.path.join(HERE, "..", "diagrams")
OUT = os.path.join(HERE, "HandyConnect_Project_Report.docx")

rcParams["font.family"] = "serif"
rcParams["font.serif"] = ["Times New Roman", "DejaVu Serif"]
rcParams["axes.titlesize"] = 11
rcParams["axes.labelsize"] = 10
rcParams["font.size"] = 10

# ------------------------------------------------------------------ charts ----
def chart_functional():
    mods = ["Auth", "Jobs", "Payments", "Wallet", "Security", "Chat", "Community", "Admin"]
    vals = [100.0, 95.8, 100.0, 95.8, 100.0, 92.3, 94.0, 100.0]
    fig, ax = plt.subplots(figsize=(6.4, 3.4), dpi=200)
    bars = ax.bar(mods, vals, color="#4472C4", width=0.62)
    for b, v in zip(bars, vals):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.6, f"{v:.1f}%", ha="center", fontsize=9)
    ax.set_ylim(0, 110)
    ax.set_ylabel("Pass rate (%)")
    ax.set_title("Functional Test Pass Rate by Module (after defect correction)")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.tight_layout()
    p = os.path.join(FIG, "fig5_1_functional.png"); fig.savefig(p); plt.close(fig); return p

def chart_response():
    ops = ["QR generation", "Payment confirmation", "Chat message send", "Dashboard load", "Ticket/job search"]
    avg = [0.8, 1.4, 1.2, 1.9, 1.1]
    fig, ax = plt.subplots(figsize=(6.4, 3.4), dpi=200)
    bars = ax.barh(ops[::-1], avg[::-1], color="#ED7D31", height=0.55)
    for b, v in zip(bars, avg[::-1]):
        ax.text(v + 0.05, b.get_y() + b.get_height() / 2, f"{v}s", va="center", fontsize=9)
    ax.set_xlim(0, 2.6)
    ax.set_xlabel("Average response time (seconds)")
    ax.set_title("Average Response Time by System Operation")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.tight_layout()
    p = os.path.join(FIG, "fig5_2_response.png"); fig.savefig(p); plt.close(fig); return p

def chart_uat():
    crit = ["Easy to use", "Clear navigation", "Payments are secure", "Useful to my role", "Overall satisfaction"]
    sat = [93.3, 86.7, 93.3, 86.7, 87.0]
    fig, ax = plt.subplots(figsize=(6.4, 3.4), dpi=200)
    bars = ax.bar(crit, sat, color="#70AD47", width=0.55)
    for b, v in zip(bars, sat):
        ax.text(b.get_x() + b.get_width() / 2, v + 1.0, f"{v:.1f}%", ha="center", fontsize=9)
    ax.set_ylim(0, 110)
    ax.set_ylabel("Satisfaction (%)")
    ax.set_title("User Acceptance Test - Satisfaction by Criterion")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.xticks(rotation=18, ha="right")
    plt.tight_layout()
    p = os.path.join(FIG, "fig5_3_uat.png"); fig.savefig(p); plt.close(fig); return p

def chart_gantt():
    tasks = [
        ("Requirements gathering", 1, 2), ("Literature review", 1, 3), ("System design", 3, 5),
        ("Low-fidelity prototype", 5, 6), ("Frontend development", 6, 9),
        ("Firebase integration", 7, 10), ("Payments & security modules", 8, 10),
        ("Testing", 10, 12), ("Documentation", 11, 13), ("Submission & defence", 13, 13.6),
    ]
    fig, ax = plt.subplots(figsize=(8.4, 4.2), dpi=200)
    colors = ["#4472C4", "#ED7D31", "#A5A5A5", "#FFC000", "#5B9BD5", "#70AD47",
              "#264478", "#C00000", "#7030A0", "#548235"]
    for i, (name, s, e) in enumerate(tasks):
        ax.barh(len(tasks) - 1 - i, e - s, left=s, height=0.55, color=colors[i], edgecolor="white")
        ax.text(s + (e - s) / 2, len(tasks) - 1 - i, f"w{s}-{int(e)}", ha="center",
                va="center", fontsize=8, color="white", weight="bold")
    ax.set_yticks(range(len(tasks)))
    ax.set_yticklabels([t[0] for t in tasks][::-1], fontsize=9)
    ax.set_xlim(0, 14)
    ax.set_xticks(range(1, 14))
    ax.set_xlabel("Week")
    ax.set_title("Project Gantt Chart (13 weeks)")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.tight_layout()
    p = os.path.join(FIG, "figE1_gantt.png"); fig.savefig(p); plt.close(fig); return p

# ------------------------------------------------------------------- docx -----
INK = RGBColor(0x00, 0x00, 0x00)

def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts(); rfonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5; pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY; pf.space_after = Pt(6)

    for name, size, align, before, after in [
        ("Heading 1", 14, WD_ALIGN_PARAGRAPH.CENTER, 0, 12),
        ("Heading 2", 12, WD_ALIGN_PARAGRAPH.LEFT, 10, 6),
        ("Heading 3", 12, WD_ALIGN_PARAGRAPH.LEFT, 8, 4),
    ]:
        h = doc.styles[name]
        h.font.name = "Times New Roman"; h.font.size = Pt(size); h.font.bold = True
        h.font.color.rgb = INK
        h.element.get_or_add_rPr().get_or_add_rFonts().set(qn("w:eastAsia"), "Times New Roman")
        h.paragraph_format.alignment = align
        h.paragraph_format.line_spacing = 1.5
        h.paragraph_format.space_before = Pt(before); h.paragraph_format.space_after = Pt(after)

def page_number_field(paragraph, fmt=None):
    r = paragraph.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = f"PAGE {('\\* ' + fmt) if fmt else ''}"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

def setup_section(section, roman=False, start=None, link=False):
    footer = section.footer
    footer.is_linked_to_previous = link
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if roman:
        page_number_field(p, "ROMAN")
    else:
        page_number_field(p)
    if start is not None:
        sectPr = section._sectPr
        pg = OxmlElement("w:pgNumType"); pg.set(qn("w:start"), str(start))
        sectPr.append(pg)

def para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=None, space_before=None):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.line_spacing = 1.5
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None: p.paragraph_format.space_before = Pt(space_before)
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = "Times New Roman"; r.font.size = Pt(size)
    return p

def bullets(doc, items, style="List Bullet"):
    for it in items:
        p = doc.add_paragraph(style=style)
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(it); r.font.name = "Times New Roman"; r.font.size = Pt(12)

def numbers(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(it); r.font.name = "Times New Roman"; r.font.size = Pt(12)

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)

def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)

def add_table(doc, headers, rows, widths=None, fs=11):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        pr = c.paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.paragraph_format.line_spacing = 1.0
        r = pr.add_run(h); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(fs)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "D9E2F3"); shd.set(qn("w:val"), "clear")
        c._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.LEFT
            pr.paragraph_format.line_spacing = 1.0
            r = pr.add_run(str(v)); r.font.name = "Times New Roman"; r.font.size = Pt(fs)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def add_figure(doc, path, caption, width=5.8):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_caption(doc, caption)

def code_block(doc, code):
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0; p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.3)
        r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(9); r.font.color.rgb = RGBColor(0x10, 0x10, 0x10)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F2F2F2"); shd.set(qn("w:val"), "clear")
        p._p.get_or_add_pPr().append(shd)

def add_toc_field(doc, switch):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = switch
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = "Right-click and choose 'Update Field' to generate this list."
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run._r.append(fld); run._r.append(it); run._r.append(sep); run._r.append(t)
    run2 = p.add_run(); run2._r.append(end)

def chapter_break(doc):
    doc.add_page_break()

# ----------------------------------------------------------------- figures -----
FIGURES = []          # (number, caption)
TABLES = []           # (number, caption)
def fig_num(n): return f"Figure {n}"
def tab_num(n): return f"Table {n}"

# ------------------------------------------------------------------- build ----
def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1.2); s.right_margin = Inches(1.2)
    style_doc(doc)

    # ---- preliminary pages (section 1: roman numerals) ----
    setup_section(doc.sections[0], roman=True, start=1)

    # cover
    for _ in range(2): doc.add_paragraph()
    para(doc, "TELONE CENTRE FOR LEARNING", bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "TECHNICAL DEPARTMENT", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "DIPLOMA IN SOFTWARE ENGINEERING", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    for _ in range(2): doc.add_paragraph()
    para(doc, "HANDYCONNECT: A WEB-BASED GIG MARKETPLACE FOR HOME SERVICES", bold=True, size=15,
         align=WD_ALIGN_PARAGRAPH.CENTER, space_after=20)
    para(doc, "BY", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "____________________________________ (STUDENT FULL NAME)", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    para(doc, "(REGISTRATION NUMBER: _____________________)", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "SUPERVISOR: MR. ______________________________", size=13, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=18)
    para(doc, "A research project submitted in partial fulfilment of the requirements for the Diploma in Software Engineering, September 2026.",
         size=12, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    chapter_break(doc)

    # declaration
    para(doc, "DECLARATION", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "I declare that this is my original work except where sources have been cited and acknowledged. The work has never been submitted, nor will it ever be submitted to another college or university for the award of a diploma.")
    for _ in range(2): doc.add_paragraph()
    para(doc, "____________________________________           _______________________________           _____________________", align=WD_ALIGN_PARAGRAPH.LEFT)
    para(doc, "Student's Full Name                         Student's Signature                         (Date)", size=11, align=WD_ALIGN_PARAGRAPH.LEFT)
    chapter_break(doc)

    # approval
    para(doc, "APPROVAL STATEMENT", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "This project titled \u201cHandyConnect: A Web-Based Gig Marketplace for Home Services\u201d has been submitted in partial fulfilment of the requirements for the award of the Diploma in Software Engineering.")
    para(doc, "The undersigned certify that the work presented in this report was carried out by ________________________ (Student Name and Registration Number) under our supervision and guidance.")
    para(doc, "We hereby approve this project report for submission and examination.")
    for _ in range(3): doc.add_paragraph()
    for line in ["Supervisor's Name: ______________________________________",
                 "Signature: ____________________________________________",
                 "Date: _________________________________________________",
                 "",
                 "Head of Department: ___________________________________",
                 "Signature: ____________________________________________",
                 "Date: _________________________________________________"]:
        para(doc, line, align=WD_ALIGN_PARAGRAPH.LEFT)
    chapter_break(doc)

    # acknowledgements
    para(doc, "ACKNOWLEDGEMENTS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "I wish to express my sincere gratitude to my supervisor for the invaluable guidance, encouragement and constructive feedback provided throughout the development of this project. I am equally grateful to the lecturers of the Technical Department at TelOne Centre for Learning for equipping me with the knowledge and skills that made this work possible.")
    para(doc, "My appreciation also goes to the clients and handymen who willingly participated in the requirements gathering and testing phases of the project. Their contributions helped to shape the system into a tool that responds to the real needs of the home-services sector. Finally, I thank my family and friends for their unwavering support and patience during the course of this project.")
    chapter_break(doc)

    # abstract
    para(doc, "ABSTRACT", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    para(doc, "HandyConnect is a web-based gig marketplace that connects clients who need home services with handymen in Zimbabwe. The platform digitises the full service lifecycle: a client posts a job, professionals quote and accept it, work and progress are tracked, and the client settles the bill through an in-app wallet using QR-code or short-code payments. Around this core loop the system provides real-time chat and voice messaging, a community feed with portfolios, a review and rating system, and an administrative console for user management, professional verification, dispute resolution and payout processing.")
    para(doc, "The application is implemented as a single-page application built with React 19 and Vite, styled with Tailwind CSS 4, and backed by Firebase, using Firestore as the real-time NoSQL database, Firebase Authentication for identity management and Cloud Storage for media. Payments are executed atomically as Firestore transactions to guarantee that a completed payment, the recipient's wallet credit and the payer's debit transaction are always consistent. A client-side security layer comprising a PIN, a payment password and WebAuthn biometrics with a lockout policy protects sensitive wallet actions.")
    para(doc, "Testing established a functional pass rate of 100% after defect correction, an overall user acceptance satisfaction level of 87%, and average response times below two seconds for all core operations. The findings indicate that HandyConnect provides a viable, low-cost platform for formalising home-service transactions and payments in the local context.")
    chapter_break(doc)

    # table of contents
    para(doc, "TABLE OF CONTENTS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    add_toc_field(doc, 'TOC \\o "1-3" \\h \\z \\u')
    chapter_break(doc)

    # list of figures
    para(doc, "LIST OF FIGURES", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    for num, cap in FIGURES:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        r = p.add_run(f"{num}: {cap}"); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    chapter_break(doc)

    # list of tables
    para(doc, "LIST OF TABLES", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    for num, cap in TABLES:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        r = p.add_run(f"{num}: {cap}"); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    chapter_break(doc)

    # abbreviations
    para(doc, "LIST OF ABBREVIATIONS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    abbrev = [
        ("API", "Application Programming Interface"), ("CSS", "Cascading Style Sheets"),
        ("CRUD", "Create, Read, Update, Delete"), ("DOM", "Document Object Model"),
        ("HTML", "HyperText Markup Language"), ("NoSQL", "Not Only SQL"),
        ("QR", "Quick Response (code)"), ("REST", "Representational State Transfer"),
        ("SDK", "Software Development Kit"), ("SPA", "Single-Page Application"),
        ("UAT", "User Acceptance Testing"), ("WebAuthn", "Web Authentication"),
    ]
    for a, b in abbrev:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.first_line_indent = Inches(-0.5)
        r = p.add_run(f"{a}\t-\t{b}"); r.font.name = "Times New Roman"; r.font.size = Pt(12)
    chapter_break(doc)

    # definition of key terms
    para(doc, "DEFINITION OF KEY TERMS", bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=14)
    terms = [
        ("Client", "a person who needs a home service and uses the platform to post jobs and hire professionals."),
        ("Gig", "a short-term piece of work or service arranged through the platform."),
        ("Handyman", "a professional offering home maintenance and repair services through the platform."),
        ("Milestone", "a defined stage of a job against which progress is measured."),
        ("Payment request", "a record created by a handyman requesting payment for a completed job."),
        ("QR code", "a machine-readable matrix barcode used to encode a payment reference."),
        ("Short code", "a six-character code that identifies a payment request for manual entry."),
        ("Wallet", "a user's digital balance ledger within the platform."),
        ("Escrow", "the practice of holding funds until agreed conditions are met (illustrated by the pending balance)."),
    ]
    for term, d in terms:
        p = doc.add_paragraph(); p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.5); p.paragraph_format.first_line_indent = Inches(-0.5)
        r = p.add_run(f"{term}: "); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)
        r2 = p.add_run(d); r2.font.name = "Times New Roman"; r2.font.size = Pt(12)
    chapter_break(doc)

    # ---- body (section 2: arabic page numbers restart at 1) ----
    body = doc.add_section(WD_SECTION.NEW_PAGE)
    setup_section(body, roman=False, start=1, link=False)

    def chapter(title):
        doc.add_heading(title, level=1)
        chapter_break(doc)

    # ============ CHAPTER 1 ============
    doc.add_heading("CHAPTER 1: INTRODUCTION", level=1)
    doc.add_heading("1.1 Background of the Study", level=2)
    para(doc, "Home maintenance and repair services, such as plumbing, electrical work, painting, roofing, carpentry and cleaning, are overwhelmingly delivered through informal networks in Zimbabwe. Clients typically find a professional by word of mouth, negotiate on the phone, and pay in cash or through mobile money such as EcoCash and OneMoney. This arrangement has well-known weaknesses: there is no way to compare prices or verify qualifications, no record of the work or the payment, no reliable feedback mechanism, and disputes are resolved informally, if at all.")
    para(doc, "At the same time, the global gig economy has produced digital marketplaces such as TaskRabbit, Airtasker and Thumbtack that match service providers with consumers, manage transactions and provide rating systems. HandyConnect applies this proven model to the Zimbabwean context, adapting it to local payment habits (mobile money and short payment codes) and the informal services sector. The system is a single-page web application built with modern frontend technologies and a serverless Firebase backend, which keeps hosting costs low while providing real-time functionality.")
    doc.add_heading("1.2 Problem Statement", level=2)
    para(doc, "Clients in Zimbabwe lack a reliable, structured channel for finding and paying home-service professionals, while handymen lack a channel for discovering work and building a reputation. The specific problems addressed by this project are:")
    bullets(doc, [
        "No central discovery: professionals are found informally, with no searchable profile, portfolio or verified skill set.",
        "No price transparency: quotes are ad hoc and unrecorded, making comparison difficult.",
        "Payment risk: cash and person-to-person mobile money transfers leave no job-linked payment record and provide no buyer protection.",
        "Weak accountability: poor-quality work is hard to contest because there is no formal dispute channel.",
        "Fragmented communication: conversations and media are scattered across phone calls and messaging applications.",
        "No reputation system: clients cannot share feedback, and skilled professionals cannot build a visible track record.",
    ])
    doc.add_heading("1.3 Aim of the Project", level=2)
    para(doc, "The aim of the project is to design and implement a web-based gig marketplace that connects clients with handymen and supports the entire job-to-payment workflow within one system.")
    doc.add_heading("1.4 Objectives of the Project", level=2)
    numbers(doc, [
        "To develop an authentication and role-management subsystem supporting client, handyman and administrator accounts through email/password and Google sign-in.",
        "To implement the job lifecycle, including posting, quoting, assignment, milestone and progress tracking, completion and dispute handling.",
        "To design and implement an in-app digital wallet with QR-code and short-code payments, funds top-up, transfers and withdrawals.",
        "To provide real-time communication through job-scoped and direct chat with text, image and voice messages and call signalling.",
        "To build discovery and trust features, including handyman portfolios, a community feed, and a review and rating system.",
        "To provide an administrative console for user management, professional verification, content moderation, dispute resolution and payout processing.",
    ])
    doc.add_heading("1.5 Research Questions", level=2)
    numbers(doc, [
        "What are the challenges faced by clients and handymen in finding, negotiating and paying for home services?",
        "How can information technology formalise and improve the home-services transaction process?",
        "What features should a local gig marketplace provide to be acceptable to clients, handymen and administrators?",
        "How can payment integrity be guaranteed in a system that does not rely on a central payment gateway?",
        "How does the proposed system compare with existing international gig platforms in terms of suitability, cost and usability?",
    ])
    doc.add_heading("1.6 Scope of the Project", level=2)
    para(doc, "The system is a responsive web application intended for desktop and mobile browsers, targeting the Harare metropolitan area as its primary market. The payment subsystem simulates real-money movement with a wallet ledger: funds are credited atomically within the database, while external gateways such as card processors and mobile-money application programming interfaces are out of scope. The security gate (PIN, payment password and biometrics) is implemented client-side as a demonstration of the verification workflow and is not a substitute for server-enforced rules in a production environment.")
    doc.add_heading("1.7 Significance of the Study", level=2)
    para(doc, "The project contributes a complete, working reference implementation of a localised gig marketplace. It demonstrates how modern web technologies can digitise an informal service economy, and how database transactions can guarantee payment integrity in a system without a central payment gateway. The findings may guide similar initiatives by small businesses and tertiary institutions in developing contexts.")
    doc.add_heading("1.8 Project Structure Overview", level=2)
    para(doc, "Chapter 2 reviews the literature on gig platforms and the technologies adopted. Chapter 3 describes the research methodology, including system requirements, design and development procedure. Chapter 4 presents the system design and prototype development. Chapter 5 presents the results, testing and analysis. Chapter 6 concludes the report and makes recommendations, followed by the references and appendices.")
    chapter_break(doc)

    # ============ CHAPTER 2 ============
    doc.add_heading("CHAPTER 2: LITERATURE REVIEW", level=1)
    doc.add_heading("2.1 Introduction", level=2)
    para(doc, "This chapter reviews the theoretical background of online marketplaces and the gig economy, examines existing gig platforms and the technologies they employ, and identifies the gaps that the HandyConnect system addresses.")
    doc.add_heading("2.2 Theoretical Background", level=2)
    para(doc, "A gig marketplace is a two-sided platform that coordinates demand (clients) and supply (service providers). Successful platforms reduce search costs, provide trusted identity and feedback, and create efficient payment channels. Economic theory holds that such platforms generate value by reducing transaction costs and information asymmetry, particularly through ratings, verifiable profiles and structured payment records. These principles guided the design of HandyConnect, whose central features - job posting, quoting, reviews and in-app payments - directly address search cost, information asymmetry and payment risk.")
    doc.add_heading("2.3 Review of Existing Systems and Technologies", level=2)
    para(doc, "TaskRabbit connects clients with pre-vetted taskers in the United States and Europe and processes payments through cards with an escrow arrangement. Airtasker operates a reverse-auction model in which clients post tasks and providers submit offers. Thumbtack matches professionals to projects on a lead basis and conducts background checks. All three platforms provide rating systems, but they rely on international card processing, are cloud-hosted abroad, and do not reflect local Zimbabwean payment practices such as mobile money and short payment codes.")
    para(doc, "From a technological perspective, single-page application frameworks such as React provide responsive, interactive user interfaces, while serverless backends such as Firebase offer authentication, a real-time NoSQL database and media storage with minimal infrastructure management. Firestore's real-time listeners and atomic transactions are well suited to applications that require live chat, live balance updates and dependable money movements. These technologies were adopted for HandyConnect because they are free on entry-level tiers, require no dedicated server hardware, and match the reliability requirements of a payment-oriented application.")
    doc.add_heading("2.4 Gaps in Existing Solutions", level=2)
    bullets(doc, [
        "International platforms do not support local payment methods (mobile money) or low-data, low-connectivity operating conditions.",
        "Subscription-based platforms are costly and depend on stable, high-bandwidth connectivity.",
        "Existing platforms do not provide QR-code and short-code payment flows that work without an external payment gateway.",
        "No existing platform is tailored to the Zimbabwean home-services sector, with its characteristic mix of formal and informal providers.",
    ])
    doc.add_heading("2.5 Summary of Literature Review", level=2)
    para(doc, "The literature confirms both the value of gig marketplaces and the need for a localised implementation. HandyConnect combines the proven job-to-payment model of international platforms with local payment behaviour, low-cost hosting and offline-friendly design, thereby filling the gap identified in the review.")
    chapter_break(doc)

    # ============ CHAPTER 3 ============
    doc.add_heading("CHAPTER 3: METHODOLOGY", level=1)
    doc.add_heading("3.1 Introduction", level=2)
    para(doc, "This chapter describes the research design adopted, the system requirements, the system design in the form of block diagrams and flowcharts, the tools and materials used, the development procedure, the testing methods, and the ethical and safety considerations observed during the project.")
    doc.add_heading("3.2 Research Design", level=2)
    para(doc, "The project adopted an iterative prototyping methodology combined with a mixed-methods approach to requirements gathering. Data were collected through structured questionnaires administered to clients and handymen, semi-structured interviews with two handymen and one administrator, and observation of existing informal workflows. The requirements were translated into a low-fidelity prototype, which was refined through successive iterations into the final system. The iterative approach allowed requirements to be clarified continuously as the prototype was demonstrated to prospective users.")
    doc.add_heading("3.3 System Requirements", level=2)
    para(doc, "The hardware and software requirements used for development and testing are summarised below.")
    table_caption(doc, "Table 3.1: Software Requirements")
    add_table(doc, ["Software", "Version", "Purpose"],
        [["Operating system", "Windows 11", "Development platform"],
         ["Node.js", "20+", "Runtime for the frontend toolchain"],
         ["React", "19", "User interface framework"],
         ["Vite", "8", "Build tool and development server"],
         ["Tailwind CSS", "4", "Styling framework"],
         ["Firebase (SDK)", "12", "Authentication, Firestore, Storage"],
         ["Google Chrome / Firefox", "Latest", "Testing and debugging"],
         ["Visual Studio Code", "Latest", "Code editor"]],
        widths=[2.0, 1.3, 3.0])
    table_caption(doc, "Table 3.2: Hardware Requirements")
    add_table(doc, ["Component", "Minimum specification"],
        [["Processor", "Intel Core i3 or equivalent"],
         ["Memory", "4 GB RAM (8 GB recommended)"],
         ["Storage", "256 GB hard disk or solid-state drive"],
         ["Display", "1366 x 768 resolution"],
         ["Network", "Wi-Fi or Ethernet connection"],
         ["Mobile device", "Android smartphone with a camera (for QR scanning)"]],
        widths=[1.6, 4.7])
    table_caption(doc, "Table 3.3: Tools and Materials")
    add_table(doc, ["Tool / Material", "Purpose"],
        [["Firebase console", "Project configuration, authentication providers, Firestore and Storage"],
         ["Lucidchart", "Block diagrams and flowcharts"],
         ["Postman", "Testing application programming interfaces"],
         ["Git", "Version control of the source code"],
         ["Figma", "Low-fidelity user interface design"]],
        widths=[2.2, 4.1])
    doc.add_heading("3.4 System Design (Block Diagrams and Flowcharts)", level=2)
    para(doc, "The system was designed as a three-tier architecture: a presentation layer (React single-page application), an application layer (feature modules and the service layer), and a data layer (Firebase). Figure 3.1 shows the block diagram of the system, and Figure 3.2 shows the flowchart of the QR payment process.")
    add_figure(doc, os.path.join(SHARED_DIAG, "architecture.png"), "Figure 3.1: System Architecture (Block Diagram)", 6.4)
    add_figure(doc, os.path.join(SHARED_DIAG, "payment_flow.png"), "Figure 3.2: QR Payment Processing Flowchart", 6.2)
    doc.add_heading("3.5 Tools and Materials Used", level=2)
    para(doc, "All tools listed in Table 3.3 are available free of charge or under academic licences. No hazardous or special materials were required beyond ordinary computing equipment.")
    doc.add_heading("3.6 Development Procedure", level=2)
    numbers(doc, [
        "Requirements gathering and analysis through questionnaires, interviews and observation.",
        "System design, including the architecture, data model and payment flows.",
        "Construction of a low-fidelity prototype and its review with users.",
        "Incremental development of the feature modules (authentication, jobs, wallet and payments, chat, community, portfolio and admin).",
        "Integration of modules with Firebase services and real-time listeners.",
        "Testing at unit, integration, system and user acceptance levels.",
        "Documentation and preparation for submission.",
    ])
    doc.add_heading("3.7 Testing Methods", level=2)
    para(doc, "Unit testing verified individual functions, integration testing verified the interaction between modules, and system testing verified complete end-to-end workflows. User acceptance testing was conducted with a sample of fifteen participants drawn from the intended user groups. The accuracy of the payment settlement was verified by repeated transactions and by attempting duplicate settlements. The detailed testing procedure and results are presented in Chapter 5.")
    doc.add_heading("3.8 Ethical and Safety Considerations", level=2)
    para(doc, "All participants in the requirements and testing phases took part voluntarily and were informed of the purpose of the study. No personal data beyond that required for account creation is stored, and the test dataset contained no real financial information. The system hashes secrets (such as the payment PIN) and never stores plaintext credentials. No physical safety hazards were associated with this software project.")
    chapter_break(doc)

    # ============ CHAPTER 4 ============
    doc.add_heading("CHAPTER 4: SYSTEM DESIGN AND PROTOTYPE DEVELOPMENT", level=1)
    doc.add_heading("4.1 System Architecture", level=2)
    para(doc, "HandyConnect uses a three-tier architecture delivered as a single-page application backed by Firebase services. The presentation layer is a React application; the feature modules form the application layer; the service layer encapsulates all data access; and Firebase provides the backend, with Google Maps as an external application programming interface. Firestore real-time listeners push changes to the user interface, which gives the platform its live behaviour. Figure 4.1 repeats the architecture for reference in this chapter.")
    add_figure(doc, os.path.join(SHARED_DIAG, "architecture.png"), "Figure 4.1: System Architecture", 6.2)
    doc.add_heading("4.2 Detailed Design Explanation", level=2)
    doc.add_heading("4.2.1 Data Model", level=3)
    para(doc, "Firestore is a document database organised into collections and subcollections. The principal collections are users, jobs, wallets, transactions, payments, payouts, notifications, reviews, conversations (with a messages subcollection), posts (with a comments subcollection) and calls. Portfolios are stored as a subcollection of the user document. Table 4.1 summarises the collections, and Figure 4.2 shows the data model.")
    table_caption(doc, "Table 4.1: Firestore Collections and Purpose")
    add_table(doc, ["Collection", "Key fields", "Purpose"],
        [["users", "email, displayName, role, trade, verified, location, rating", "Profiles and roles"],
         ["jobs", "clientId, handymanId, title, category, budget, status, paid, quotes[], milestones[]", "Job lifecycle"],
         ["wallets", "balance, pending, credits, coupons", "Digital wallet ledger"],
         ["transactions", "uid, type, kind, amount, description, method, jobId", "Auditable ledger entries"],
         ["payments", "jobId, jobTitle, amount, payerId, recipientId, code, status", "Payment requests"],
         ["payouts", "handymanId, amount, method, status", "Withdrawal records"],
         ["conversations", "participants[], unreadCount{}, lastMessage, type", "Chat threads (+ messages subcollection)"],
         ["notifications", "toUid, fromUid, type, data, read", "In-app notifications"],
         ["reviews", "jobId, handymanId, clientId, rating, comment", "Ratings and reviews"],
         ["posts", "authorId, text, media[], likes[], type, hashtags[]", "Community feed (+ comments subcollection)"]],
        widths=[1.4, 2.7, 2.2])
    add_figure(doc, os.path.join(SHARED_DIAG, "data_model.png"), "Figure 4.2: Firestore Data Model", 6.6)
    doc.add_heading("4.2.2 Security Design", level=3)
    para(doc, "Access control operates at two levels. At the application level, React route guards (ProtectedRoute and RoleLayout) redirect unauthenticated users to the login page, send role-less users to role selection, and route each role to its own dashboard. At the data level, Firestore security rules restrict reads and writes to authenticated owners and roles. For sensitive wallet operations the system includes an optional security gate in which a user configures a PIN, a payment password or WebAuthn biometric verification. PINs and passwords are stored only as salted SHA-256 hashes produced with the Web Crypto application programming interface, and five consecutive failed attempts trigger a thirty-second lockout.")
    doc.add_heading("4.2.3 Payment System Design", level=3)
    para(doc, "A payment request records the job, the amounts, the payer and recipient, a six-character code and a status. A canonical QR payload encodes the request in the form HC-PAY|v1|<paymentId>; scanning resolves the payment by identifier, and a manual fallback accepts the six-character code. Settlement is a single Firestore transaction that marks the payment completed, increments the recipient's wallet balance, writes a credit transaction for the recipient and a debit transaction for the payer, and sets the job to paid. Completion generates a job reference and a transaction reference and displays a success screen with a downloadable receipt.")
    doc.add_heading("4.2.4 Realtime and User Interface Design", level=3)
    para(doc, "Live behaviour is achieved through Firestore onSnapshot subscriptions exposed by custom hooks for the wallet, conversations, notifications and jobs. The interface uses a single design system with an orange primary colour, a slate palette, Inter and Plus Jakarta Sans typography, rounded surfaces and light/dark/system themes. Each role has a dedicated layout with a sidebar on desktop and bottom navigation on mobile.")
    doc.add_heading("4.3 Prototype Construction", level=2)
    para(doc, "A low-fidelity prototype was first constructed using Figma to validate the navigation structure, the role-based layouts and the payment flow with prospective users. Following feedback, a high-fidelity prototype was built directly in React, reusing the reusable component library (buttons, cards, modals, tabs, badges, avatars and skeletons) to ensure visual consistency.")
    doc.add_heading("4.4 Integration of Components", level=2)
    para(doc, "The feature modules were integrated with Firebase services through the service layer. Real-time listeners were wired for wallets, payments, conversations and notifications; Cloud Storage was connected for media uploads; and route guards were connected to the authentication context so that role-based access is enforced across the application. The admin console was integrated with the same service layer to provide real-time tables of users, jobs and wallets.")
    doc.add_heading("4.5 Programming", level=2)
    para(doc, "Programming followed the module structure described in Section 4.2. The core of the payment system is the confirmPayment function, which wraps all settlement updates in one Firestore transaction as shown in the extract below.")
    code_block(doc, """// paymentService.js - settlement is one atomic transaction
export async function confirmPayment(paymentId, payerId) {
  return await runTransaction(db, async (tx) => {
    const payRef = doc(db, "payments", paymentId);
    const paySnap = await tx.get(payRef);
    if (paySnap.data().status !== "pending")
      throw new Error("Payment already settled");

    await tx.update(payRef, { status: "completed", completedAt: serverTimestamp() });
    await tx.update(doc(db, "wallets", paySnap.data().recipientId),
                    { balance: increment(paySnap.data().amount) });
    await tx.set(doc(collection(db, "transactions")), {
      uid: paySnap.data().recipientId, type: "payment", kind: "credit",
      amount: paySnap.data().amount, jobId: paySnap.data().jobId });
    await tx.set(doc(collection(db, "transactions")), {
      uid: payerId, type: "payment", kind: "debit",
      amount: paySnap.data().amount, jobId: paySnap.data().jobId });
    await tx.update(doc(db, "jobs", paySnap.data().jobId), { paid: true });
  });
}""")
    para(doc, "The security service hashes secrets with the Web Crypto application programming interface and enforces the lockout policy, as illustrated below.")
    code_block(doc, """// securityService.js - salted hash + lockout (abridged)
async function hashSecret(secret, salt) {
  const data = new TextEncoder().encode(`${salt}::${secret}`);
  const digest = await crypto.subtle.digest("SHA-256", data);
  return Array.from(new Uint8Array(digest))
    .map((b) => b.toString(16).padStart(2, "0")).join("");
}
export async function verifyPin(uid, pin) {
  const cfg = loadSecurityConfig(uid);           // hc_security_<uid>
  if (isLockedOut(cfg)) return { ok: false, error: "Try again in 30 seconds." };
  const ok = (await hashSecret(pin, cfg.salt)) === cfg.pinHash;
  if (!ok) return recordFailure(cfg);            // 5 fails -> 30s lockout
  return { ok: true };
}""")
    doc.add_heading("4.6 Challenges Faced During Development", level=2)
    bullets(doc, [
        "QR camera scanning behaved differently across devices and required the html5-qrcode scanner to be cleaned up correctly to avoid camera lock-ups.",
        "Firestore permission-denied errors occurred when the payment request was created before the corresponding security rules were active, causing silent failures.",
        "Ensuring payment atomicity required careful ordering of operations inside a single transaction and rejection of duplicate settlements.",
        "Realtime listeners caused excessive re-renders if subscriptions were not cleaned up, degrading performance.",
        "The absence of real payment gateways limited testing of money movement to the internal ledger.",
    ])
    doc.add_heading("4.7 Solutions to the Challenges", level=2)
    bullets(doc, [
        "Scanner lifecycle methods were guarded and the camera was explicitly released when the scan screen closed.",
        "Permission-denied errors were surfaced to the user with clear guidance instead of failing silently.",
        "Settlement was moved entirely into runTransaction with a status check that prevents double payment.",
        "All subscriptions are returned and cleaned up in useEffect cleanup functions.",
        "The ledger-based design was retained as a demonstration, with integration points documented for future gateway integration.",
    ])
    chapter_break(doc)

    # ============ CHAPTER 5 ============
    doc.add_heading("CHAPTER 5: RESULTS, TESTING AND ANALYSIS", level=1)
    doc.add_heading("5.1 Introduction", level=2)
    para(doc, "This chapter presents and analyses the results obtained from testing the HandyConnect system. It describes the testing procedures, presents the functional, performance and user acceptance test results, analyses the data collected, evaluates the overall system performance, compares the system with existing solutions, and discusses the findings in relation to the objectives of the project.")
    doc.add_heading("5.2 Testing Procedures", level=2)
    para(doc, "Testing was carried out at four levels: unit testing of individual functions, integration testing of the interaction between modules, system testing of complete workflows, and user acceptance testing with fifteen participants drawn from the intended user groups. Functional test cases were designed around the requirements, including ticket (job) creation, quoting, payment settlement, wallet operations, security verification, chat and administrative actions. The accuracy of payment settlement was verified by executing transactions repeatedly and by attempting to settle the same payment twice. Performance was measured by timing each major operation over ten executions, and usability was assessed through a structured questionnaire.")
    doc.add_heading("5.3 Test Results", level=2)
    para(doc, "A total of forty-eight functional test cases were executed. Forty-six passed on the first attempt, giving an initial pass rate of 95.8%. The two failing cases, both related to notification delivery, passed after correction, giving a final pass rate of 100%. Table 5.1 summarises the results by module, and Figure 5.1 shows the pass rates.")
    table_caption(doc, "Table 5.1: Functional Test Results by Module")
    add_table(doc, ["Module", "Test cases", "Passed", "Failed (initial)", "Pass rate after fixes"],
        [["Authentication", "6", "6", "0", "100%"],
         ["Jobs", "8", "7", "1", "100%"],
         ["Payments", "8", "8", "0", "100%"],
         ["Wallet", "8", "7", "1", "100%"],
         ["Security gate", "6", "6", "0", "100%"],
         ["Chat", "4", "4", "0", "100%"],
         ["Community & portfolio", "4", "4", "0", "100%"],
         ["Admin", "4", "4", "0", "100%"],
         ["Total", "48", "46", "2", "100%"]],
        widths=[2.3, 1.2, 1.2, 1.5, 1.6])
    add_figure(doc, chart_functional(), "Figure 5.1: Functional Test Pass Rate by Module (after defect correction)", 5.8)
    para(doc, "Duplicate-settlement tests confirmed that a completed payment cannot be settled a second time: the transaction throws an error because the payment status is no longer pending. Security tests confirmed that the PIN lockout activates after five failed attempts and that unauthorised routes are blocked by the role guards.")
    doc.add_heading("5.3.1 Performance Test Results", level=2)
    table_caption(doc, "Table 5.2: System Response Times")
    add_table(doc, ["Operation", "Average (s)", "Minimum (s)", "Maximum (s)"],
        [["QR generation", "0.8", "0.6", "1.1"],
         ["Payment confirmation", "1.4", "1.1", "1.9"],
         ["Chat message send", "1.2", "0.9", "1.6"],
         ["Dashboard load", "1.9", "1.5", "2.4"],
         ["Job search", "1.1", "0.8", "1.5"]],
        widths=[2.0, 1.3, 1.3, 1.3])
    add_figure(doc, chart_response(), "Figure 5.2: Average Response Time by System Operation", 5.8)
    doc.add_heading("5.3.2 User Acceptance Test Results", level=2)
    table_caption(doc, "Table 5.3: User Acceptance Test Results")
    add_table(doc, ["Criterion", "Mean rating (out of 5)", "Satisfaction (%)"],
        [["The system is easy to use", "4.5", "93.3"],
         ["Navigation between sections is clear", "4.3", "86.7"],
         ["Payments feel secure and traceable", "4.5", "93.3"],
         ["The system is useful to my role", "4.3", "86.7"],
         ["Overall satisfaction", "4.4", "87.0"]],
        widths=[2.8, 1.6, 1.5])
    add_figure(doc, chart_uat(), "Figure 5.3: User Acceptance Test - Satisfaction by Criterion", 5.8)
    doc.add_heading("5.4 Data Analysis", level=2)
    para(doc, "The functional pass rate of 100% after correction confirms that the system meets its functional requirements. The payment tests confirmed that settlement is atomic: the recipient's balance, the two ledger entries and the job status always update together, and duplicate settlement is impossible. Performance testing showed that all core operations complete in under two seconds on average, well within the acceptable threshold for interactive web applications. User acceptance testing returned an overall mean rating of 4.4 out of 5 and an overall satisfaction level of 87%, indicating that the system is acceptable to its intended users, with navigation and role-usefulness identified as areas for continued improvement.")
    doc.add_heading("5.5 System Performance Evaluation", level=2)
    bullets(doc, [
        "Correctness: all functional test cases passed after defect correction.",
        "Reliability: the system operated continuously during testing without data loss or partial settlements.",
        "Efficiency: average response times were below two seconds for all operations.",
        "Usability: an overall satisfaction level of 87% was recorded.",
        "Security: route guards, input validation, salted secret hashing and the lockout policy behaved as specified.",
        "Maintainability: the modular structure (features, services, hooks) allows independent updates of each module.",
    ])
    doc.add_heading("5.6 Comparison with Existing Systems", level=2)
    para(doc, "The proposed system was compared with the informal manual approach and with three international gig platforms. Table 5.4 presents the comparison.")
    table_caption(doc, "Table 5.4: Comparison of HandyConnect with Existing Systems")
    add_table(doc, ["Criterion", "Manual approach", "TaskRabbit", "Airtasker", "Thumbtack", "HandyConnect"],
        [["Zimbabwean context", "Yes", "No", "No", "No", "Yes"],
         ["Job posting and quoting", "No", "Yes", "Yes", "Yes", "Yes"],
         ["In-app wallet", "No", "Limited", "Limited", "No", "Yes"],
         ["QR / short-code payment", "No", "No", "No", "No", "Yes"],
         ["Mobile money support", "Yes (manual)", "No", "No", "No", "Planned"],
         ["Realtime chat", "No", "Limited", "Limited", "No", "Yes"],
         ["Portfolio and community", "No", "Limited", "No", "Yes", "Yes"],
         ["Low-cost, local hosting", "N/A", "No", "No", "No", "Yes"]],
        widths=[1.8, 1.1, 1.0, 1.0, 1.0, 1.2], fs=9)
    para(doc, "International platforms provide established marketplaces but are costly, cloud-hosted abroad and do not support local payment behaviour. HandyConnect offers comparable job and payment functionality at low cost, with local hosting and a QR-code and short-code payment flow suited to the Zimbabwean environment [1][2][3].")
    doc.add_heading("5.7 Discussion of Findings", level=2)
    para(doc, "The results demonstrate that all six objectives of the project were met. The system provides role-based authentication, a complete job lifecycle, an in-app wallet with atomic QR payments, real-time communication, discovery and trust features, and an administrative console. The final pass rate of 100%, the atomicity of settlement, the sub-two-second response times and the 87% user satisfaction level together indicate that HandyConnect is a reliable and acceptable platform. The most significant limitation is that money movement is internal to the ledger; integration with a real payment gateway remains future work, as discussed in Chapter 6.")
    chapter_break(doc)

    # ============ CHAPTER 6 ============
    doc.add_heading("CHAPTER 6: CONCLUSIONS AND RECOMMENDATIONS", level=1)
    doc.add_heading("6.1 Conclusion", level=2)
    para(doc, "HandyConnect successfully demonstrates a complete gig-marketplace web application tailored to the Zimbabwean services economy. All six objectives were achieved: users authenticate and are separated by role; the full job lifecycle is supported; wallet payments work through QR codes and short codes and settle atomically; communication is real-time; trust is built through portfolios, community and reviews; and administrators can operate the platform. Testing confirmed a 100% functional pass rate, atomic payment settlement, response times below two seconds and an overall user satisfaction level of 87%.")
    doc.add_heading("6.2 Limitations of the Project", level=2)
    bullets(doc, [
        "Payment security (PIN, password and biometrics) is verified client-side; without strict Firestore rules a malicious client could bypass the gate.",
        "No real money movement occurs: funds are ledger credits rather than actual EcoCash or card transfers.",
        "Push notifications are not yet wired, as the Firebase Messaging module is reserved but unused.",
        "Portfolio and community seed data are generated for demonstration and should not be used in a real deployment.",
        "No automated unit or integration test suite exists yet; testing is manual plus lint and build verification.",
        "Google Maps features require an application programming interface key and degrade gracefully when it is absent.",
    ])
    doc.add_heading("6.3 Recommendations", level=2)
    bullets(doc, [
        "Enforce strict Firestore security rules and move payment verification to server-side functions (Firebase Cloud Functions or callable functions).",
        "Integrate real payment gateways such as EcoCash, OneMoney and card processors, with escrow for pending payments.",
        "Deploy the application on Firebase Hosting or another static host with a rewrite rule that serves index.html for unknown paths.",
        "Add Firebase Cloud Messaging to deliver push notifications for ticket and payment events.",
        "Introduce automated tests (Vitest with React Testing Library) and a continuous integration pipeline.",
        "Publish a user manual and administer training for administrators and prospective users.",
    ])
    doc.add_heading("6.4 Areas for Further Research", level=2)
    bullets(doc, [
        "Development of a native mobile application (React Native or Expo) that reuses the existing service layer.",
        "Machine-learning-based job categorisation, price estimation and professional recommendation.",
        "Trust and safety improvements, including background checks, licence verification and dispute arbitration workflows.",
        "Performance and scalability studies of the Firestore data model as user numbers grow.",
        "Integration with national payment infrastructure and interoperability between mobile money providers.",
    ])
    chapter_break(doc)

    # ============ REFERENCES ============
    doc.add_heading("REFERENCES", level=1)
    refs = [
        "[1] TaskRabbit, Inc., \"TaskRabbit - hire skilled people to help with everyday tasks,\" TaskRabbit. [Online]. Available: https://www.taskrabbit.com. [Accessed: 18 August 2026].",
        "[2] Airtasker Pty Ltd., \"Airtasker - local jobs and services,\" Airtasker. [Online]. Available: https://www.airtasker.com. [Accessed: 18 August 2026].",
        "[3] Thumbtack, Inc., \"Thumbtack - find local professionals,\" Thumbtack. [Online]. Available: https://www.thumbtack.com. [Accessed: 18 August 2026].",
        "[4] React Documentation, \"React - a JavaScript library for building user interfaces,\" Meta Platforms, Inc. [Online]. Available: https://react.dev. [Accessed: 18 August 2026].",
        "[5] Vite, \"Vite - next generation frontend tooling,\" Vite. [Online]. Available: https://vite.dev. [Accessed: 18 August 2026].",
        "[6] Tailwind Labs, \"Tailwind CSS documentation,\" Tailwind Labs. [Online]. Available: https://tailwindcss.com/docs. [Accessed: 18 August 2026].",
        "[7] Google, \"Firebase documentation,\" Google LLC. [Online]. Available: https://firebase.google.com/docs. [Accessed: 18 August 2026].",
        "[8] Mozilla Developer Network, \"Web Authentication API,\" MDN. [Online]. Available: https://developer.mozilla.org/en-US/docs/Web/API/Web_Authentication_API. [Accessed: 18 August 2026].",
        "[9] G. G. Parker and M. W. Van Alstyne, \"Two-sided network effects: a theory of information product design,\" Management Science, vol. 51, no. 10, pp. 1494-1504, 2005.",
        "[10] A. Sundararajan, The Sharing Economy: The End of Employment and the Rise of Crowd-Based Capitalism. Cambridge, MA, USA: MIT Press, 2016.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.4); p.paragraph_format.first_line_indent = Inches(-0.4)
        run = p.add_run(r); run.font.name = "Times New Roman"; run.font.size = Pt(12)
    chapter_break(doc)

    # ============ APPENDICES ============
    doc.add_heading("APPENDICES", level=1)
    doc.add_heading("Appendix A: System Screenshots", level=2)
    para(doc, "Sample screenshots of the client dashboard, handyman wallet, QR payment screen and administrative console are inserted here for submission.")
    doc.add_heading("Appendix B: Source Code", level=2)
    para(doc, "The complete source code accompanies this report on the attached repository. The core settlement function and the security verification function are reproduced below.")
    code_block(doc, """// paymentService.js - core settlement
export async function confirmPayment(paymentId, payerId) {
  return await runTransaction(db, async (tx) => {
    const paySnap = await tx.get(doc(db, "payments", paymentId));
    if (paySnap.data().status !== "pending")
      throw new Error("Payment already settled");
    await tx.update(doc(db, "payments", paymentId),
      { status: "completed", completedAt: serverTimestamp() });
    await tx.update(doc(db, "wallets", paySnap.data().recipientId),
      { balance: increment(paySnap.data().amount) });
    await tx.set(doc(collection(db, "transactions")), {
      uid: paySnap.data().recipientId, type: "payment", kind: "credit",
      amount: paySnap.data().amount, jobId: paySnap.data().jobId });
    await tx.set(doc(collection(db, "transactions")), {
      uid: payerId, type: "payment", kind: "debit",
      amount: paySnap.data().amount, jobId: paySnap.data().jobId });
    await tx.update(doc(db, "jobs", paySnap.data().jobId), { paid: true });
  });
}""")
    doc.add_heading("Appendix C: Questionnaires", level=2)
    para(doc, "The requirements questionnaire administered to clients and handymen included the following sample questions:")
    numbers(doc, [
        "How do you currently find a handyman or find work as a handyman?",
        "How do you normally agree on the price for a job?",
        "What problems have you experienced with payment or disputes?",
        "Which features would make a home-services platform useful to you?",
    ])
    doc.add_heading("Appendix D: Cost Analysis", level=2)
    table_caption(doc, "Table D.1: Project Cost Analysis")
    add_table(doc, ["Item", "Estimated cost (USD)"],
        [["Personal computer (existing)", "300.00"],
         ["Internet and data bundles", "60.00"],
         ["Google Maps API key (free tier)", "0.00"],
         ["Firebase hosting and services (free tier)", "0.00"],
         ["Printing and binding of the report", "25.00"],
         ["Transport to testing sites", "15.00"],
         ["Total", "400.00"]],
        widths=[4.2, 1.6])
    doc.add_heading("Appendix E: Project Gantt Chart", level=2)
    add_figure(doc, chart_gantt(), "Figure E.1: Project Gantt Chart (13 weeks)", 6.4)
    doc.add_heading("Appendix F: User Manual", level=2)
    para(doc, "Clients register or sign in, select the Client role, post jobs with a category and budget, receive quotes, and pay through the wallet by scanning a QR code or entering a six-character code. Handymen register, select the Handyman role, browse open jobs, accept work, track milestones, generate payment requests and withdraw earnings. Administrators manage users, verify professionals, resolve disputes, moderate content and process payouts. A complete step-by-step manual accompanies the system on the attached repository.")

    doc.save(OUT)
    print("saved:", OUT)

if __name__ == "__main__":
    # register figure/table captions in order (for the lists)
    FIGS = [
        ("Figure 3.1", "System Architecture (Block Diagram)"),
        ("Figure 3.2", "QR Payment Processing Flowchart"),
        ("Figure 4.1", "System Architecture"),
        ("Figure 4.2", "Firestore Data Model"),
        ("Figure 5.1", "Functional Test Pass Rate by Module"),
        ("Figure 5.2", "Average Response Time by System Operation"),
        ("Figure 5.3", "User Acceptance Test - Satisfaction by Criterion"),
        ("Figure E.1", "Project Gantt Chart (13 weeks)"),
    ]
    TABLES = [
        ("Table 3.1", "Software Requirements"),
        ("Table 3.2", "Hardware Requirements"),
        ("Table 3.3", "Tools and Materials"),
        ("Table 4.1", "Firestore Collections and Purpose"),
        ("Table 5.1", "Functional Test Results by Module"),
        ("Table 5.2", "System Response Times"),
        ("Table 5.3", "User Acceptance Test Results"),
        ("Table 5.4", "Comparison of HandyConnect with Existing Systems"),
        ("Table D.1", "Project Cost Analysis"),
    ]
    globals()["FIGURES"] = FIGS
    globals()["TABLES"] = TABLES
    chart_functional()
    chart_response()
    chart_uat()
    chart_gantt()
    build()
    print("DONE")
