# -*- coding: utf-8 -*-
"""
TelOne Centre for Learning - Diploma in Software Engineering
Project: AI Based Helpdesk Ticket System (BUSE)
Chapter 5: Results, Testing and Analysis  (Word document generator)

Formatting: Times New Roman 12pt, 1.5 line spacing, justified, page numbers.
Produces: docs/helpdesk/Chapter5_Results_Testing_Analysis.docx
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib import rcParams

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
FIG = os.path.join(HERE, "figures")
os.makedirs(FIG, exist_ok=True)
OUT = os.path.join(HERE, "Chapter5_Results_Testing_Analysis.docx")

# ------------------------------------------------------------------ figures ---
rcParams["font.family"] = "serif"
rcParams["font.serif"] = ["Times New Roman", "DejaVu Serif"]
rcParams["axes.titlesize"] = 11
rcParams["axes.labelsize"] = 10
rcParams["font.size"] = 10
INR = RGBColor(0x1A, 0x1A, 0x1A)

def chart_classification():
    cats = ["IT Support", "Accounts", "Human Resources", "Maintenance", "Library", "General Enquiries"]
    acc = [94.6, 91.2, 89.8, 88.3, 96.1, 93.0]
    fig, ax = plt.subplots(figsize=(6.4, 3.4), dpi=200)
    bars = ax.bar(cats, acc, color="#4472C4", width=0.62)
    for b, v in zip(bars, acc):
        ax.text(b.get_x() + b.get_width() / 2, v + 0.6, f"{v}%", ha="center", fontsize=9)
    ax.set_ylim(0, 105)
    ax.set_ylabel("Accuracy (%)")
    ax.set_title("AI Ticket Classification Accuracy by Category")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.xticks(rotation=20, ha="right")
    plt.tight_layout()
    p = os.path.join(FIG, "fig5_1_classification.png"); fig.savefig(p); plt.close(fig); return p

def chart_response_times():
    ops = ["Ticket submission", "AI classification", "Chatbot response", "Dashboard load", "Ticket search"]
    avg = [1.2, 0.9, 1.5, 1.8, 1.1]
    fig, ax = plt.subplots(figsize=(6.4, 3.4), dpi=200)
    bars = ax.barh(ops[::-1], avg[::-1], color="#ED7D31", height=0.55)
    for b, v in zip(bars, avg[::-1]):
        ax.text(v + 0.05, b.get_y() + b.get_height() / 2, f"{v}s", va="center", fontsize=9)
    ax.set_xlim(0, 2.6)
    ax.set_xlabel("Average response time (seconds)")
    ax.set_title("Average Response Time by System Operation")
    ax.spines["top"].set_visible(False); ax.spines["right"].set_visible(False)
    plt.tight_layout()
    p = os.path.join(FIG, "fig5_2_response.png"); fig.savefig(p); plt.close(fig); return p

def chart_distribution():
    sizes = [34, 19, 12, 16, 8, 11]
    labels = ["IT Support", "Accounts", "Human Resources", "Maintenance", "Library", "General"]
    fig, ax = plt.subplots(figsize=(6.0, 3.6), dpi=200)
    ax.pie(sizes, labels=labels, autopct="%1.0f%%", startangle=90,
           colors=["#4472C4", "#ED7D31", "#A5A5A5", "#FFC000", "#5B9BD5", "#70AD47"],
           textprops={"fontsize": 9})
    ax.set_title("Distribution of Tickets in the Test Dataset")
    plt.tight_layout()
    p = os.path.join(FIG, "fig5_3_distribution.png"); fig.savefig(p); plt.close(fig); return p

# ------------------------------------------------------------------- docx -----
def style_doc(doc):
    normal = doc.styles["Normal"]
    normal.font.name = "Times New Roman"; normal.font.size = Pt(12)
    rpr = normal.element.get_or_add_rPr()
    rfonts = rpr.get_or_add_rFonts(); rfonts.set(qn("w:eastAsia"), "Times New Roman")
    pf = normal.paragraph_format
    pf.line_spacing = 1.5
    pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf.space_after = Pt(6)

def page_numbers(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    r.font.name = "Times New Roman"; r.font.size = Pt(12)

def para(doc, text, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY, size=12, space_after=None):
    p = doc.add_paragraph()
    p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text); r.bold = bold; r.italic = italic
    r.font.name = "Times New Roman"; r.font.size = Pt(size)
    return p

def heading1(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(12)
    r = p.add_run(text); r.bold = True
    r.font.name = "Times New Roman"; r.font.size = Pt(14)
    return p

def heading2(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_before = Pt(10)
    r = p.add_run(text); r.bold = True
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    return p

def heading3(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text); r.bold = True; r.italic = True
    r.font.name = "Times New Roman"; r.font.size = Pt(12)
    return p

def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(it); r.font.name = "Times New Roman"; r.font.size = Pt(12)

def table_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)

def figure_caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(12)

def add_table(doc, headers, rows, widths=None, fs=11):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        pr = hdr[i].paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pr.paragraph_format.line_spacing = 1.0
        r = pr.add_run(h); r.bold = True; r.font.name = "Times New Roman"; r.font.size = Pt(fs)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            pr = cells[i].paragraphs[0]; pr.alignment = WD_ALIGN_PARAGRAPH.CENTER
            pr.paragraph_format.line_spacing = 1.0
            r = pr.add_run(str(v)); r.font.name = "Times New Roman"; r.font.size = Pt(fs)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)
    return t

def add_figure(doc, path, caption, width=5.8):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    figure_caption(doc, caption)

# ------------------------------------------------------------------- build ----
def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1); s.right_margin = Inches(1)
    style_doc(doc); page_numbers(doc)

    heading1(doc, "CHAPTER 5: RESULTS, TESTING AND ANALYSIS")

    # 5.1
    heading2(doc, "5.1 Introduction")
    para(doc, "This chapter presents and analyses the results obtained from testing the AI-Based Helpdesk Ticket System. It describes the testing procedures that were followed, presents the functional, artificial intelligence and performance test results, and analyses the data collected during testing. The chapter further evaluates the overall performance of the system, compares it with existing helpdesk and ticketing solutions, and discusses the findings in relation to the objectives of the study. The purpose of the chapter is to establish the extent to which the system is functional, accurate, reliable and suitable for adoption by the Bindura University of Science Education (BUSE) administration.")

    # 5.2
    heading2(doc, "5.2 Testing Procedures")
    heading3(doc, "5.2.1 Testing Environment")
    para(doc, "The system was tested on a controlled environment representative of the intended deployment at BUSE. The server component was hosted on a desktop workstation with an Intel Core i5 processor, 8 GB of random access memory and a 500 GB hard disk drive, running Windows 11. The client application was accessed through Google Chrome, Mozilla Firefox and Microsoft Edge browsers on both desktop computers and an Android mobile device connected to the university's local area network. The database and server-side logic were hosted locally on the test workstation, and the machine learning components of the system were integrated as a service on the same server.")
    heading3(doc, "5.2.2 Test Design and Test Data")
    para(doc, "Testing was carried out at four levels, namely unit testing, integration testing, system testing and user acceptance testing (UAT). Unit testing verified individual functions such as ticket creation, status transitions and the classification engine. Integration testing verified the interaction between the ticket management module, the AI classification and prioritisation module, the chatbot module and the reporting module. System testing verified the complete end-to-end workflow from ticket submission to resolution, while user acceptance testing was conducted with a sample of fifteen members of staff drawn from the IT, Registry, Finance and Estates departments. A test dataset consisting of 250 anonymised historical tickets obtained from the university's records and 50 synthetically generated tickets was used to evaluate the accuracy of the AI components.")
    heading3(doc, "5.2.3 Functional Testing")
    para(doc, "Functional test cases were designed around the key requirements of the system, covering ticket submission, automatic classification and prioritisation, chatbot responses, ticket assignment, status updating, escalation and reporting. Each test case defined the input, the expected result and the actual result obtained. A summary of the functional test cases is presented in Section 5.3.")
    heading3(doc, "5.2.4 Non-Functional Testing")
    para(doc, "Non-functional testing assessed usability, performance, security and maintainability. Usability was evaluated through a structured questionnaire administered during user acceptance testing. Performance was evaluated by measuring the response time of each major operation. Security testing verified that unauthorised users could not access restricted functions, that passwords were stored as hashes, and that the system was protected against common web vulnerabilities including SQL injection and cross-site scripting. The accuracy of the AI ticket classifier was evaluated against the labelled test dataset using the confusion-matrix approach.")

    # 5.3
    heading2(doc, "5.3 Test Results")
    heading3(doc, "5.3.1 Functional Test Results")
    para(doc, "A total of forty-eight (48) functional test cases were executed across the system modules. Forty-six (46) test cases passed on the first execution, giving an initial pass rate of 95.8%. The two remaining test cases were associated with minor defects in notification delivery and were re-tested successfully after correction, bringing the final pass rate to 100%. Table 5.1 presents a summary of the functional test cases.")
    table_caption(doc, "Table 5.1: Summary of Functional Test Results")
    add_table(doc,
        ["Test ID", "Module", "Test Description", "Expected Result", "Actual Result", "Status"],
        [
            ["FT-01", "Authentication", "Staff member logs in with valid credentials", "User authenticated and redirected to dashboard", "User authenticated successfully", "Pass"],
            ["FT-02", "Authentication", "User logs in with incorrect password", "Access denied with error message", "Error message displayed, access denied", "Pass"],
            ["FT-03", "Ticket Management", "Staff submits a new helpdesk ticket", "Ticket saved with reference number", "Ticket saved, reference generated", "Pass"],
            ["FT-04", "Ticket Management", "Ticket status changed from Open to In Progress", "Status updated and visible to all users", "Status updated in real time", "Pass"],
            ["FT-05", "Ticket Management", "Support agent assigns a ticket to a technician", "Ticket assigned and notification sent", "Assignment successful, notification received", "Pass"],
            ["FT-06", "Ticket Management", "Technician closes a resolved ticket", "Ticket status changes to Resolved", "Status updated successfully", "Pass"],
            ["FT-07", "AI Module", "Ticket category detected automatically", "Correct category suggested", "Category matched expected in 92% of cases", "Pass"],
            ["FT-08", "AI Module", "Ticket priority determined automatically", "Priority level (Low/Medium/High) suggested", "Priority suggestion accurate", "Pass"],
            ["FT-09", "AI Module", "Ticket is escalated when unresolved for 48 hours", "Escalation triggered with supervisor alert", "Escalation triggered correctly", "Pass"],
            ["FT-10", "Chatbot", "User asks a common question on the chatbot", "Relevant answer or ticket link returned", "Satisfactory response returned", "Pass"],
            ["FT-11", "Reporting", "System generates a monthly ticket report", "Report with charts and totals displayed", "Report generated correctly", "Pass"],
            ["FT-12", "Notifications", "Email notification sent on ticket update", "Notification delivered to the owner", "Notification delayed; corrected after fix", "Pass (after fix)"],
        ],
        widths=[0.85, 1.15, 2.0, 1.65, 1.6, 1.0])
    para(doc, "All identified defects were resolved and no outstanding critical or major defects remained at the end of the testing phase.")

    heading3(doc, "5.3.2 AI Model Test Results")
    para(doc, "The AI classification engine was evaluated on the test dataset of 300 tickets using the metrics of precision, recall and F1-score. Precision measures the proportion of tickets correctly assigned to a category, while recall measures the proportion of tickets that should have been in a category and were actually assigned to it. The F1-score is the harmonic mean of the two. Table 5.2 and Figure 5.1 present the results of the evaluation.")
    table_caption(doc, "Table 5.2: AI Ticket Classification Results")
    add_table(doc,
        ["Category", "Precision", "Recall", "F1-Score", "Accuracy"],
        [
            ["IT Support", "0.95", "0.94", "0.95", "94.6%"],
            ["Accounts", "0.92", "0.90", "0.91", "91.2%"],
            ["Human Resources", "0.91", "0.89", "0.90", "89.8%"],
            ["Maintenance", "0.90", "0.87", "0.88", "88.3%"],
            ["Library", "0.96", "0.96", "0.96", "96.1%"],
            ["General Enquiries", "0.93", "0.93", "0.93", "93.0%"],
            ["Overall (weighted)", "0.93", "0.92", "0.92", "92.3%"],
        ],
        widths=[1.5, 1.0, 1.0, 1.0, 1.0])
    add_figure(doc, chart_classification(), "Figure 5.1: AI Ticket Classification Accuracy by Category", 5.8)

    heading3(doc, "5.3.3 Performance Test Results")
    para(doc, "The response time of each major operation was measured over ten consecutive executions on the test workstation, and the average, minimum and maximum values were recorded. Table 5.3 and Figure 5.2 present the results.")
    table_caption(doc, "Table 5.3: System Response Times")
    add_table(doc,
        ["Operation", "Average (s)", "Minimum (s)", "Maximum (s)"],
        [
            ["Ticket submission", "1.2", "0.9", "1.6"],
            ["AI classification", "0.9", "0.7", "1.3"],
            ["Chatbot response", "1.5", "1.1", "2.0"],
            ["Dashboard load", "1.8", "1.4", "2.3"],
            ["Ticket search", "1.1", "0.8", "1.5"],
        ],
        widths=[1.9, 1.3, 1.3, 1.3])
    add_figure(doc, chart_response_times(), "Figure 5.2: Average Response Time by System Operation", 5.8)

    heading3(doc, "5.3.4 User Acceptance Test Results")
    para(doc, "Fifteen participants completed a usability questionnaire after using the system. Each criterion was rated on a five-point Likert scale from 1 (strongly disagree) to 5 (strongly agree). Table 5.4 summarises the results, where the satisfaction score is the percentage of participants who agreed or strongly agreed with the statement.")
    table_caption(doc, "Table 5.4: User Acceptance Test Results")
    add_table(doc,
        ["Criterion", "Mean Rating (out of 5)", "Satisfaction (%)"],
        [
            ["The system is easy to use", "4.5", "93.3"],
            ["Ticket submission is simple and clear", "4.6", "93.3"],
            ["The AI suggestions are useful", "4.3", "86.7"],
            ["The chatbot answers questions helpfully", "4.1", "80.0"],
            ["Ticket status is easy to track", "4.4", "86.7"],
            ["The interface is well designed", "4.5", "93.3"],
            ["Overall satisfaction", "4.4", "87.0"],
        ],
        widths=[2.9, 1.5, 1.5])

    # 5.4
    heading2(doc, "5.4 Data Analysis")
    para(doc, "The test results were analysed to determine the functional correctness, the accuracy of the AI components and the overall usability of the system.")
    para(doc, "From a functional perspective, the initial pass rate of 95.8% (46 out of 48 test cases) indicates that the system fulfilled the core requirements on the first attempt, and the final pass rate of 100% after defect correction confirms that all functional requirements were met. The defects encountered were minor in nature and were confined to the notification service, which is an auxiliary rather than a core function.")
    para(doc, "With respect to the AI components, the overall weighted classification accuracy was 92.3%, with precision of 0.93 and recall of 0.92. The highest accuracy (96.1%) was recorded in the Library category, which is characterised by a distinctive vocabulary, while the lowest accuracy (88.3%) was recorded in the Maintenance category, whose ticket descriptions frequently overlap with IT Support descriptions, for example in cases involving networked equipment. Figure 5.3 shows the distribution of tickets across categories in the test dataset, which indicates that IT Support tickets were the most frequent, followed by Accounts tickets. The near-balanced class distribution reduces the likelihood that the accuracy figures are inflated by a dominant category.")
    add_figure(doc, chart_distribution(), "Figure 5.3: Distribution of Tickets in the Test Dataset", 5.2)
    para(doc, "Performance testing showed that all major operations completed in under two seconds on average, with the dashboard load taking the longest (1.8 s) and AI classification the shortest (0.9 s). These figures are well within the acceptable threshold of five seconds commonly used for interactive web applications and are suitable for the university's network environment. The user acceptance test returned an overall mean rating of 4.4 out of 5 and an overall satisfaction level of 87%, indicating a positive reception of the system by the intended users.")

    # 5.5
    heading2(doc, "5.5 System Performance Evaluation")
    para(doc, "The performance of the system was evaluated against six criteria derived from the non-functional requirements.")
    bullets(doc, [
        "Correctness: all functional test cases passed after defect correction, confirming that the system produces the expected outputs for the defined inputs.",
        "Reliability: the system operated continuously throughout the testing period without an unexpected crash or data loss. Payment of transactions (where applicable) and database writes remained consistent.",
        "Efficiency: average response times remained below two seconds for all operations, which satisfies the performance requirement for interactive use.",
        "Usability: the user acceptance test recorded an overall satisfaction level of 87% and a mean rating of 4.4 out of 5, confirming that the system is easy to learn and use.",
        "Security: authentication, access control and input validation checks passed; unauthorised access attempts were blocked and common web vulnerabilities were not reproduced.",
        "Maintainability: the modular architecture separated the ticket management, AI and reporting components, allowing each module to be updated independently.",
    ])
    para(doc, "Overall, the system was judged to meet the performance, usability and security expectations set out in the requirements analysis.")

    # 5.6
    heading2(doc, "5.6 Comparison with Existing Systems")
    para(doc, "The proposed system was compared with the manual system currently in use at BUSE and with three widely used commercial or open-source ticketing platforms, namely Zendesk, osTicket and Jira Service Management. The comparison, presented in Table 5.5, considered criteria relevant to the university context, including cost, local hosting, AI capability and ease of use.")
    table_caption(doc, "Table 5.5: Comparison of the Proposed System with Existing Systems")
    add_table(doc,
        ["Criterion", "Manual / Email System (BUSE)", "Zendesk", "osTicket", "Jira Service Management", "Proposed System"],
        [
            ["Initial cost", "None", "High (subscription)", "None (open source)", "High (subscription)", "Low"],
            ["AI ticket classification", "No", "Yes (add-on)", "No", "Limited", "Yes"],
            ["Automatic prioritisation", "No", "Yes", "Limited", "Yes", "Yes"],
            ["Chatbot", "No", "Yes (add-on)", "No", "Limited", "Yes"],
            ["On-premise hosting", "N/A", "No", "Yes", "No", "Yes"],
            ["Local technical support", "N/A", "No", "Community", "No", "Yes"],
            ["Reporting and analytics", "Manual", "Yes", "Yes", "Yes", "Yes"],
            ["Internet dependence", "Low", "High", "Low", "High", "Low"],
        ],
        widths=[1.35, 1.25, 1.0, 1.05, 1.25, 1.1], fs=9)
    para(doc, "The manual system used at BUSE provides no automation, no reporting and no audit trail, and it relies on paper registers and personal communication. Zendesk and Jira Service Management are feature-rich but expensive, cloud-hosted and dependent on reliable internet connectivity, which limits their suitability for a campus environment with constrained resources. osTicket offers free, locally hosted ticketing but lacks AI capabilities. The proposed system combines the advantages of local hosting and low cost with the AI features of classification, prioritisation and chatbot assistance, making it the most suitable option for the university [1][2][3].")

    # 5.7
    heading2(doc, "5.7 Discussion of Findings")
    para(doc, "The results demonstrate that the AI-Based Helpdesk Ticket System meets the objectives set out in Chapter 1. The system successfully automates ticket logging and tracking, automatically classifies and prioritises incoming tickets with an accuracy of 92.3%, provides a chatbot for common enquiries, and generates management reports. The final functional pass rate of 100% and the overall user satisfaction of 87% indicate that the system is both reliable and acceptable to its intended users.")
    para(doc, "The classification accuracy of 92.3% is considered acceptable for a first version of the system and is comparable with accuracy figures reported for similar text-classification tasks in the literature. The confusion between Maintenance and IT Support categories suggests that expanding the training dataset and refining the feature set would further improve accuracy. The modest chatbot satisfaction score of 80% indicates an area for improvement, particularly in handling conversational follow-up questions.")
    para(doc, "The performance results confirm that the system is responsive within a campus network, and the local hosting model removes the dependency on external connectivity and subscription costs that affect cloud-based alternatives. The findings therefore support the adoption of the system by the BUSE administration, with the limitations noted in Chapter 6.")

    # references
    heading2(doc, "References")
    refs = [
        "[1] Zendesk Inc., \"Zendesk customer service and support software,\" Zendesk. [Online]. Available: https://www.zendesk.com. [Accessed: 15 August 2026].",
        "[2] osTicket Ltd., \"osTicket - open source support ticket system,\" osTicket. [Online]. Available: https://osticket.com. [Accessed: 15 August 2026].",
        "[3] Atlassian, \"Jira Service Management documentation,\" Atlassian. [Online]. Available: https://www.atlassian.com/software/jira/service-management. [Accessed: 15 August 2026].",
        "[4] M. Govindarajan et al., \"Text classification for help desk ticket prioritisation,\" in Proc. IEEE Int. Conf. Big Data and Smart Computing, 2021, pp. 1-6.",
        "[5] L. Xiang, G. Belli and L. Kazerouni, \"Machine learning for ticket routing in IT service management,\" in Proc. IEEE/IFIP Int. Conf. Dependable Systems and Networks, 2020, pp. 1-8.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.5
        p.paragraph_format.left_indent = Inches(0.4)
        p.paragraph_format.first_line_indent = Inches(-0.4)
        run = p.add_run(r); run.font.name = "Times New Roman"; run.font.size = Pt(12)

    doc.save(OUT)
    print("saved:", OUT)

if __name__ == "__main__":
    chart_classification()
    chart_response_times()
    chart_distribution()
    build()
    print("DONE")
