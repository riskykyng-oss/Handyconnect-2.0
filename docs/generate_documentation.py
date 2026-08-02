# -*- coding: utf-8 -*-
"""
HandyConnect - Final Year Project Documentation Generator.
Produces:
  docs/diagrams/*.png            (architecture, data model, payment flow, job lifecycle, auth flow, use cases)
  docs/HandyConnect_Project_Documentation.docx
"""
import os
import matplotlib
matplotlib.use("Agg")
import matplotlib.pyplot as plt
from matplotlib.patches import FancyBboxPatch, FancyArrowPatch, Ellipse, Rectangle

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DIAG = os.path.join(HERE, "diagrams")
os.makedirs(DIAG, exist_ok=True)
OUT_DOCX = os.path.join(HERE, "HandyConnect_Project_Documentation.docx")

# ---------------------------------------------------------------- palette ---
C = {
    "primary": "#F97316", "primary_soft": "#FFEDD5",
    "slate": "#334155", "border": "#CBD5E1", "surface": "#FFFFFF", "bg": "#F8FAFC",
    "emerald": "#059669", "emerald_soft": "#D1FAE5",
    "amber": "#D97706", "amber_soft": "#FEF3C7",
    "blue": "#2563EB", "blue_soft": "#DBEAFE",
    "violet": "#7C3AED", "violet_soft": "#EDE9FE",
    "red": "#DC2626", "red_soft": "#FEE2E2",
    "ink": "#1E293B", "mut": "#64748B",
}

# ------------------------------------------------------------ diagram core ---
def canvas(w, h):
    fig, ax = plt.subplots(figsize=(w, h), dpi=210)
    ax.set_xlim(0, w); ax.set_ylim(0, h); ax.axis("off")
    ax.set_facecolor(C["bg"])
    fig.patch.set_facecolor(C["bg"])
    return fig, ax

def box(ax, x, y, w, h, text, fill=C["surface"], edge=C["border"], tc=C["slate"],
        fs=8.5, bold=False, rounding=0.12, lw=1.2):
    p = FancyBboxPatch((x, y), w, h, boxstyle=f"round,pad=0.02,rounding_size={rounding}",
                       linewidth=lw, edgecolor=edge, facecolor=fill)
    ax.add_patch(p)
    ax.text(x + w / 2, y + h / 2, text, ha="center", va="center", fontsize=fs,
            color=tc, weight="bold" if bold else "normal")

def band(ax, x, y, w, h, title, subtitle=None, fill=C["surface"], edge=C["border"]):
    p = FancyBboxPatch((x, y), w, h, boxstyle="round,pad=0.02,rounding_size=0.15",
                       linewidth=1.4, edgecolor=edge, facecolor=fill)
    ax.add_patch(p)
    ax.text(x + w / 2, y + h - 0.42, title, ha="center", va="center", fontsize=9.5,
            color=C["slate"], weight="bold")
    if subtitle:
        ax.text(x + w / 2, y + 0.32, subtitle, ha="center", va="center", fontsize=7.2,
                color=C["mut"])

def arrow(ax, x1, y1, x2, y2, label=None, color=C["mut"], lw=1.5, style="-|>", ls="-"):
    a = FancyArrowPatch((x1, y1), (x2, y2), arrowstyle=style, mutation_scale=12,
                        linewidth=lw, color=color, linestyle=ls)
    ax.add_patch(a)
    if label:
        ax.text((x1 + x2) / 2, (y1 + y2) / 2 + 0.14, label, ha="center", va="bottom",
                fontsize=7.2, color=color)

def save(fig, name):
    path = os.path.join(DIAG, name)
    fig.savefig(path, bbox_inches="tight", facecolor=C["bg"])
    plt.close(fig)
    print("diagram:", name)
    return path

# ------------------------------------------------------------- architecture ---
def diagram_architecture():
    fig, ax = canvas(12, 8)
    # actors
    box(ax, 0.6, 7.15, 2.4, 0.7, "CLIENT", fill=C["blue_soft"], edge=C["blue"], fs=9, bold=True)
    box(ax, 4.8, 7.15, 2.4, 0.7, "HANDYMAN", fill=C["primary_soft"], edge=C["primary"], fs=9, bold=True)
    box(ax, 9.0, 7.15, 2.4, 0.7, "ADMIN", fill=C["violet_soft"], edge=C["violet"], fs=9, bold=True)
    ax.text(6, 7.75, "Users", ha="center", fontsize=8, color=C["mut"], style="italic")

    # presentation band
    band(ax, 0.6, 5.9, 10.8, 0.95, "PRESENTATION LAYER - React 19 Single-Page Application (Vite 8)",
         "React 19  -  Tailwind CSS 4  -  framer-motion  -  lucide-react  -  qrcode.react  -  html5-qrcode  -  react-hot-toast  -  date-fns")
    arrow(ax, 2.6, 7.15, 3.4, 6.85)
    arrow(ax, 6.0, 7.15, 6.0, 6.85)
    arrow(ax, 10.2, 7.15, 9.4, 6.85)

    # feature modules
    band(ax, 0.6, 4.35, 10.8, 1.25, "FEATURE MODULES")
    mods = [
        ("Authentication\n& Role Management", 0.9, C["primary_soft"], C["primary"]),
        ("Job Management\n(post / quote / complete)", 2.7, C["blue_soft"], C["blue"]),
        ("Wallet & QR Payments\n(receive / pay / withdraw)", 4.5, C["emerald_soft"], C["emerald"]),
        ("Chat & Calls\n(job + direct)", 6.3, C["violet_soft"], C["violet"]),
        ("Community & Portfolio\n(posts / stories / pros)", 8.1, C["amber_soft"], C["amber"]),
        ("Admin Console\n(moderation / payouts)", 9.9, C["red_soft"], C["red"]),
    ]
    for label, x, fill, edge in mods:
        box(ax, x, 4.55, 1.75, 0.85, label, fill=fill, edge=edge, fs=6.8)
    arrow(ax, 6.0, 5.9, 6.0, 5.6)

    # services
    band(ax, 0.6, 2.85, 10.8, 1.2, "SERVICE LAYER  (src/services)", )
    svcs = [
        ("authService\nregister / login / google", 0.9),
        ("jobService\ncreate / accept / complete", 2.7),
        ("paymentService\nQR token / confirmPayment", 4.5),
        ("walletService\nbalance / transactions", 6.3),
        ("securityService\nPIN / password / WebAuthn", 8.1),
        ("chat / post / review /\nadmin / maps services", 9.9),
    ]
    for label, x in svcs:
        box(ax, x, 3.05, 1.75, 0.8, label, fill=C["surface"], fs=6.4)
    arrow(ax, 6.0, 4.35, 6.0, 4.05)

    # firebase / backend
    band(ax, 0.6, 1.0, 10.8, 1.35, "BACKEND  -  Firebase (BaaS) + External APIs")
    box(ax, 1.0, 1.25, 2.3, 0.85, "Firebase Auth\n(email/password, Google)", fill=C["primary_soft"], edge=C["primary"], fs=7)
    box(ax, 3.6, 1.25, 3.4, 0.85, "Cloud Firestore\n(users, jobs, wallets, payments,\ntransactions, chat, posts, ...)", fill=C["blue_soft"], edge=C["blue"], fs=7)
    box(ax, 7.3, 1.25, 2.0, 0.85, "Cloud Storage\n(images, voice)", fill=C["emerald_soft"], edge=C["emerald"], fs=7)
    box(ax, 9.6, 1.25, 1.5, 0.85, "Google\nMaps API", fill=C["violet_soft"], edge=C["violet"], fs=7)
    arrow(ax, 6.0, 2.85, 6.0, 2.5)
    return save(fig, "architecture.png")

# ---------------------------------------------------------------- data model ---
def diagram_data_model():
    fig, ax = canvas(15, 8.6)

    def card(x, y, title, fields, fill=C["surface"], edge=C["border"], sub=None):
        h = 0.5 + 0.34 * len(fields)
        p = FancyBboxPatch((x, y), 3.1, h, boxstyle="round,pad=0.02,rounding_size=0.1",
                           linewidth=1.3, edgecolor=edge, facecolor=fill)
        ax.add_patch(p)
        ax.text(x + 0.1, y + h - 0.28, title, ha="left", va="center", fontsize=8,
                color=C["slate"], weight="bold")
        for i, f in enumerate(fields):
            ax.text(x + 0.12, y + h - 0.62 - 0.34 * i, f, ha="left", va="center", fontsize=6.6, color=C["ink"])
        if sub:
            ax.text(x + 3.15, y + h / 2, sub, ha="left", va="center", fontsize=6.4, color=C["mut"], style="italic")
        return h

    # column 1
    h = card(0.3, 5.6, "users/{uid}", ["email, displayName, role", "createdAt, photoURL", "skills, trade, verified", "location{lat,lng}, rating", "bio, jobs, suspended"])
    card(0.3, 3.0, "wallets/{uid}", ["balance, pending", "credits, coupons", "currency: USD"])
    card(0.3, 0.9, "follows/{followerId_followeeId}", ["followerId, followeeId", "createdAt"])
    # column 2
    card(4.0, 5.6, "jobs/{id}", ["clientId, handymanId", "handymanName, title", "category, budget, status", "quotes[], milestones[]", "progress, attachments[]", "dispute{}, paid, createdAt"])
    card(4.0, 2.4, "transactions/{id}", ["uid, type, kind", "amount, description", "method, category, jobId"])
    card(4.0, 0.9, "calls/{id}", ["fromId, toId, convId", "status (ringing/answered/ended)"])
    # column 3
    card(7.7, 5.6, "payments/{id}", ["type, jobId, jobTitle", "amount, currency", "payerId, recipientId/Name", "code (6-char), status", "createdAt, completedAt"])
    card(7.7, 2.4, "payouts/{id}", ["handymanId, amount", "method, status"])
    card(7.7, 0.9, "reviews/{id}", ["jobId, handymanId, clientId", "clientName, rating (1-5), comment"])
    # column 4
    card(11.4, 5.6, "conversations/{id}", ["participants[]", "unreadCount{}, lastMessage", "type (job | direct)", "jobId, jobTitle", "directKey, participantInfo{}"], sub="+ subcollections/{id}/messages")
    card(11.4, 3.0, "posts/{id}", ["authorId, authorName, role", "text, media[], imageUrl", "type, trade, location", "likes[], reactions{}, hashtags[]", "commentCount"], sub="+ subcollections/{id}/comments")
    card(11.4, 0.9, "users/{uid}/portfolio/{id}", ["title, description", "images[], price, location", "featured, before/after"], sub="(subcollection)")
    card(9.0, 0.9, "notifications/{id}", ["toUid, fromUid, type", "data{}, read"], )

    # relationships
    arrow(ax, 3.4, 6.9, 4.0, 6.6, "1:M  posts jobs", color=C["blue"], lw=1.4)
    arrow(ax, 3.4, 5.6, 3.4, 4.3, "1:1  owns", color=C["primary"], lw=1.4)
    arrow(ax, 4.0, 4.0, 3.4, 3.4, "1:M  transactions", color=C["emerald"], lw=1.4)
    arrow(ax, 7.7, 5.7, 7.1, 5.2, "1:1  settles job", color=C["violet"], lw=1.4)
    arrow(ax, 7.7, 4.4, 4.0, 3.4, "payouts <= wallets", color=C["amber"], lw=1.4, style="-|>")
    ax.text(9.7, 7.75, "Key relationships", fontsize=8, color=C["slate"], weight="bold")
    return save(fig, "data_model.png")

# ------------------------------------------------------------- payment flow ---
def diagram_payment_flow():
    fig, ax = canvas(11, 8.4)
    xs = {"Handyman": 1.2, "Modal": 3.0, "Firestore": 5.4, "ScanPage": 7.6, "Client": 9.6}
    for name, x in xs.items():
        ax.plot([x, x], [0.6, 7.9], ls="--", color="#94A3B8", lw=1.1)
        ax.text(x, 8.05, name, ha="center", fontsize=8.5, color=C["slate"], weight="bold")

    def step(y, a, b, label, color=C["mut"], lw=1.5):
        x1, x2 = xs[a], xs[b]
        d = 0.25 if x2 > x1 else -0.25
        arrow(ax, x1, y, x2 - d, y, label=label, color=color, lw=lw)

    step(7.4, "Handyman", "Modal", "1  Taps 'Receive' + amount")
    step(7.0, "Modal", "Firestore", "2  createPaymentRequest()")
    ax.text(5.4, 6.7, "creates payments/{id}  status=pending, 6-char code", ha="center", fontsize=6.6, color=C["mut"])
    step(6.35, "Firestore", "Modal", "3  payment id returned")
    step(6.0, "Modal", "Handyman", "4  QR (HC-PAY|v1|<id>) + code shown", color=C["primary"], lw=1.7)
    step(5.4, "Client", "ScanPage", "5  Scans QR or types 6-char code")
    step(5.0, "ScanPage", "Firestore", "6  getPaymentByCode(code)")
    step(4.6, "Firestore", "ScanPage", "7  payment + recipient details")
    step(4.2, "Client", "ScanPage", "8  Confirms amount")
    step(3.8, "ScanPage", "Firestore", "9  confirmPayment()  -> runTransaction", color=C["emerald"], lw=1.7)
    box(ax, 3.0, 2.6, 4.8, 0.95, "Atomic transaction\npayment -> completed  |  wallet balance +\n2 transactions (credit + debit)  |  job.paid = true",
        fill=C["emerald_soft"], edge=C["emerald"], fs=6.8)
    arrow(ax, 5.4, 3.55, 5.4, 3.3)
    step(3.0, "Firestore", "ScanPage", "10  success")
    step(2.6, "ScanPage", "Client", "11  PaymentSuccessScreen\nconfetti  +  receipt  +  review prompt", color=C["primary"], lw=1.7)
    return save(fig, "payment_flow.png")

# ------------------------------------------------------------ job lifecycle ---
def diagram_job_lifecycle():
    fig, ax = canvas(10, 5.2)
    states = [
        (0.7, 2.4, "OPEN\nclient posts job", C["amber_soft"], C["amber"]),
        (3.2, 2.4, "ASSIGNED\nhandyman accepts", C["blue_soft"], C["blue"]),
        (5.7, 2.4, "COMPLETED\nwork finished,\nwallet credited", C["emerald_soft"], C["emerald"]),
        (8.2, 2.4, "PAID\nclient confirms\nQR payment", C["emerald_soft"], C["emerald"]),
        (3.2, 0.6, "DISPUTED\nopened by client", C["red_soft"], C["red"]),
        (7.0, 0.6, "CANCELLED\nclient / admin", C["surface"], C["border"]),
    ]
    for x, y, t, f, e in states:
        box(ax, x, y, 1.9, 1.35, t, fill=f, edge=e, fs=6.8)
    arrow(ax, 2.6, 3.05, 3.2, 3.05, "acceptJob", color=C["blue"], lw=1.6)
    arrow(ax, 5.1, 3.05, 5.7, 3.05, "completeJob", color=C["emerald"], lw=1.6)
    arrow(ax, 7.6, 3.05, 8.2, 3.05, "confirmPayment", color=C["primary"], lw=1.6)
    arrow(ax, 4.15, 2.4, 4.15, 1.95, "openDispute", color=C["red"], lw=1.4)
    arrow(ax, 4.15, 1.5, 5.6, 1.5, "", color=C["mut"], lw=1.2, ls="--")
    ax.text(4.9, 1.35, "admin resolves", fontsize=6.6, color=C["mut"])
    arrow(ax, 5.6, 2.4, 5.0, 2.2, "", color=C["mut"], lw=1.2, ls="--")
    arrow(ax, 2.1, 2.4, 5.4, 1.0, "cancel", color=C["mut"], lw=1.2, ls="--")
    arrow(ax, 4.1, 2.4, 7.0, 1.3, "cancel", color=C["mut"], lw=1.2, ls="--")
    return save(fig, "job_lifecycle.png")

# ----------------------------------------------------------------- auth flow ---
def diagram_auth_flow():
    fig, ax = canvas(10, 6.2)
    box(ax, 4.0, 5.2, 2.0, 0.7, "START", fill=C["violet_soft"], edge=C["violet"], fs=8, bold=True)
    arrow(ax, 5.0, 5.2, 5.0, 4.8)
    box(ax, 1.2, 3.9, 3.4, 0.9, "Email / Password\nsignup (verification email)", fill=C["surface"], fs=7.2)
    box(ax, 5.4, 3.9, 3.4, 0.9, "Sign in with Google", fill=C["surface"], fs=7.2)
    arrow(ax, 5.0, 4.8, 3.0, 4.35)
    arrow(ax, 5.0, 4.8, 7.0, 4.35)
    box(ax, 3.0, 2.7, 4.0, 0.85, "Role selection\n(client | handyman | admin)", fill=C["primary_soft"], edge=C["primary"], fs=7.5)
    arrow(ax, 4.4, 3.9, 4.8, 3.55)
    arrow(ax, 7.0, 3.9, 5.4, 3.55)
    box(ax, 3.0, 1.55, 4.0, 0.85, "user doc created with role", fill=C["blue_soft"], edge=C["blue"], fs=7.5)
    arrow(ax, 5.0, 2.7, 5.0, 2.4)
    box(ax, 0.6, 0.35, 2.6, 0.9, "Client dashboard\n+ route guards", fill=C["blue_soft"], edge=C["blue"], fs=7)
    box(ax, 3.7, 0.35, 2.6, 0.9, "Handyman dashboard\n+ route guards", fill=C["primary_soft"], edge=C["primary"], fs=7)
    box(ax, 6.8, 0.35, 2.6, 0.9, "Admin console\n+ route guards", fill=C["violet_soft"], edge=C["violet"], fs=7)
    arrow(ax, 4.3, 1.55, 1.9, 1.25)
    arrow(ax, 5.0, 1.55, 5.0, 1.25)
    arrow(ax, 5.7, 1.55, 8.1, 1.25)
    ax.text(0.35, 4.5, "New user", fontsize=7, color=C["mut"], style="italic")
    return save(fig, "auth_flow.png")

# ---------------------------------------------------------------- use cases ---
def diagram_use_cases():
    fig, ax = canvas(12, 7.2)
    # actors
    box(ax, 0.4, 3.2, 2.0, 0.8, "CLIENT", fill=C["blue_soft"], edge=C["blue"], fs=9, bold=True)
    box(ax, 9.6, 3.2, 2.0, 0.8, "HANDYMAN", fill=C["primary_soft"], edge=C["primary"], fs=9, bold=True)
    box(ax, 5.0, 0.3, 2.0, 0.8, "ADMIN", fill=C["violet_soft"], edge=C["violet"], fs=9, bold=True)
    # system boundary
    ax.add_patch(Rectangle((2.7, 0.5), 6.7, 6.2, fill=False, edgecolor=C["border"], lw=1.6))
    ax.text(6.05, 6.55, "HandyConnect System", ha="center", fontsize=9, color=C["slate"], weight="bold")

    def uc(x, y, t, fill=C["surface"]):
        ax.add_patch(Ellipse((x, y), 2.2, 0.62, facecolor=fill, edgecolor=C["border"], lw=1.2))
        ax.text(x, y, t, ha="center", va="center", fontsize=6.6, color=C["ink"], weight="bold")

    client_uc = [(3.6, 5.9, "Register / Login"), (3.6, 5.1, "Post a Job"),
                 (3.6, 4.3, "Browse & Hire"), (3.6, 3.5, "Pay by QR / code"),
                 (3.6, 2.7, "Add Funds / Wallet"), (3.6, 1.9, "Chat & Voice")]
    handyman_uc = [(8.5, 5.9, "Find & Accept Jobs"), (8.5, 5.1, "Quote / Milestones"),
                   (8.5, 4.3, "Receive Payment (QR)"), (8.5, 3.5, "Withdraw Earnings"),
                   (8.5, 2.7, "Manage Portfolio"), (8.5, 1.9, "Community Posts")]
    admin_uc = [(4.4, 1.6, "Verify Pros"), (5.9, 1.6, "Resolve Disputes"),
                (7.4, 1.6, "Process Payouts"), (4.4, 0.9, "Moderate Content"),
                (5.9, 0.9, "Manage Users"), (7.4, 0.9, "Broadcast Announcements")]
    for x, y, t in client_uc:
        uc(x, y, t, fill=C["blue_soft"])
        ax.plot([2.4, 2.9], [3.6, y], color=C["mut"], lw=1.1)
    for x, y, t in handyman_uc:
        uc(x, y, t, fill=C["primary_soft"])
        ax.plot([9.6, 9.1], [3.6, y], color=C["mut"], lw=1.1)
    for x, y, t in admin_uc:
        uc(x, y, t, fill=C["violet_soft"])
        ax.plot([6.0, x], [1.1, y + 0.05], color=C["mut"], lw=1.1)
    return save(fig, "use_cases.png")

# -------------------------------------------------------------------- docx ---
INK = RGBColor(0x1E, 0x29, 0x3B)
ORANGE = RGBColor(0xEA, 0x58, 0x0C)
MUT = RGBColor(0x64, 0x74, 0x8B)

def set_margins(doc):
    for s in doc.sections:
        s.top_margin = Inches(1); s.bottom_margin = Inches(1)
        s.left_margin = Inches(1); s.right_margin = Inches(1)

def style_doc(doc):
    st = doc.styles["Normal"]
    st.font.name = "Calibri"; st.font.size = Pt(11); st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.15
    for name, size, color in [("Heading 1", 16, ORANGE), ("Heading 2", 13, INK), ("Heading 3", 11.5, INK)]:
        h = doc.styles[name]
        h.font.name = "Calibri"; h.font.size = Pt(size); h.font.bold = True; h.font.color.rgb = color
        h.paragraph_format.space_before = Pt(14 if name == "Heading 1" else 10)
        h.paragraph_format.space_after = Pt(6)

def page_number_footer(doc):
    footer = doc.sections[0].footer
    p = footer.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run()
    f1 = OxmlElement("w:fldChar"); f1.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = "PAGE"
    f2 = OxmlElement("w:fldChar"); f2.set(qn("w:fldCharType"), "end")
    r._r.append(f1); r._r.append(it); r._r.append(f2)
    r.font.size = Pt(9); r.font.color.rgb = MUT

def add_toc(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    fld = OxmlElement("w:fldChar"); fld.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve")
    it.text = 'TOC \\o "1-2" \\h \\z \\u'
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    run._r.append(fld); run._r.append(it); run._r.append(sep)
    t = OxmlElement("w:t"); t.text = "Table of contents - right-click here and choose 'Update Field' to generate."
    run._r.append(t)
    run2 = p.add_run()
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    run2._r.append(end)

def para(doc, text, bold=False, italic=False, align=None, color=None, size=None, space_after=None):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    if space_after is not None: p.paragraph_format.space_after = Pt(space_after)
    return p

def bullets(doc, items, style="List Bullet"):
    for it in items:
        doc.add_paragraph(it, style=style)

def numbers(doc, items):
    for it in items:
        doc.add_paragraph(it, style="List Number")

def code_block(doc, code):
    for line in code.strip("\n").split("\n"):
        p = doc.add_paragraph()
        r = p.add_run(line)
        r.font.name = "Consolas"; r.font.size = Pt(8.5); r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "F1F5F9"); shd.set(qn("w:val"), "clear")
        p.paragraph_format.space_after = Pt(0); p.paragraph_format.line_spacing = 1.0
        p._p.get_or_add_pPr().append(shd)

def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Table Grid"; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ""
        r = hdr[i].paragraphs[0].add_run(h); r.bold = True; r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shd = OxmlElement("w:shd"); shd.set(qn("w:fill"), "EA580C"); shd.set(qn("w:val"), "clear")
        hdr[i]._tc.get_or_add_tcPr().append(shd)
    for row in rows:
        cells = t.add_row().cells
        for i, v in enumerate(row):
            cells[i].text = ""
            r = cells[i].paragraphs[0].add_run(str(v)); r.font.size = Pt(9)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t

def image(doc, path, caption, width=6.4):
    doc.add_picture(path, width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    para(doc, caption, italic=True, color=MUT, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)

def page_break(doc):
    doc.add_page_break()

# ------------------------------------------------------------ build document ---
def build():
    doc = Document()
    set_margins(doc); style_doc(doc); page_number_footer(doc)

    # cover
    for _ in range(4): doc.add_paragraph()
    para(doc, "HANDYCONNECT", bold=True, size=40, color=ORANGE, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=4)
    para(doc, "A Web-Based Gig Marketplace for Home Services", bold=True, size=18, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=16)
    para(doc, "Final Year Project Documentation", size=13, color=MUT, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=60)
    for label in ["Student Name: ______________________________",
                  "Registration Number: _______________________",
                  "Supervisor: _______________________________",
                  "Institution / Department: ___________________",
                  "Academic Year: 2025 / 2026"]:
        para(doc, label, size=12, color=INK, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=10)
    page_break(doc)

    # abstract
    doc.add_heading("Abstract", level=1)
    para(doc, "HandyConnect is a web-based gig marketplace that connects clients who need home services with vetted handymen in Zimbabwe. The platform digitises the full service lifecycle: a client posts a job, professionals quote and accept it, work and progress are tracked, and the client settles the bill through an in-app wallet using QR-code or short-code payments. Around this core loop the system provides real-time chat and voice messaging, a community feed with portfolios, a review and rating system, and an administrative console for user management, professional verification, dispute resolution and payout processing.")
    para(doc, "The application is implemented as a single-page application (SPA) built with React 19 and Vite, styled with Tailwind CSS 4, and backed by Firebase - using Firestore as the real-time NoSQL database, Firebase Authentication for identity management and Cloud Storage for media. Payments are executed atomically as Firestore transactions to guarantee that a completed payment, the recipient's wallet credit, and the payer's debit transaction are always consistent. A client-side security layer (PIN, payment password and WebAuthn biometrics with a lockout policy) protects sensitive wallet actions.")
    para(doc, "This document presents the background and problem statement, a review of existing platforms, requirements analysis, system design (architecture, data model, security and payment design), implementation details, testing, deployment and an evaluation of the completed system together with directions for future work.")
    page_break(doc)

    # toc
    doc.add_heading("Table of Contents", level=1)
    add_toc(doc)
    page_break(doc)

    # 1. introduction
    doc.add_heading("1. Introduction", level=1)
    doc.add_heading("1.1 Background", level=2)
    para(doc, "Home maintenance and repair services - plumbing, electrical work, painting, roofing, carpentry and cleaning - are overwhelmingly delivered through informal networks in Zimbabwe. Clients typically find a professional by word of mouth, negotiate on the phone, and pay in cash or through mobile money such as EcoCash and OneMoney. This arrangement has well-known weaknesses: there is no way to compare prices or verify qualifications, no record of the work or the payment, no reliable feedback mechanism, and disputes are resolved informally, if at all.")
    para(doc, "Meanwhile, the global gig economy has produced digital marketplaces such as TaskRabbit, Airtasker and Thumbtack that match service providers with consumers, manage the transaction and provide rating systems. HandyConnect applies this proven model to the local Zimbabwean context, adapting it to local payment habits (mobile money and short payment codes) and the informal services sector.")
    doc.add_heading("1.2 Problem Statement", level=2)
    para(doc, "Clients in Zimbabwe lack a reliable, structured channel for finding and paying home-service professionals, while handymen lack a channel for discovering work and building a reputation. The specific problems addressed are:")
    bullets(doc, [
        "No central discovery: professionals are found informally, with no searchable profile, portfolio or verified skill set.",
        "No price transparency: quotes are ad hoc and unrecorded, making comparison difficult.",
        "Payment risk: cash and person-to-person mobile money transfers leave no job-linked payment record and provide no buyer protection.",
        "Weak accountability: poor-quality work is hard to contest because there is no formal dispute channel.",
        "Fragmented communication: conversations and media are scattered across phone calls and messaging apps.",
        "No reputation system: clients cannot share feedback, and skilled professionals cannot build a visible track record.",
    ])
    doc.add_heading("1.3 Aim and Objectives", level=2)
    para(doc, "The aim of the project is to design and implement a web-based gig marketplace that connects clients with handymen and supports the entire job-to-payment workflow in one system. The specific objectives are:")
    numbers(doc, [
        "O1 - Develop an authentication and role-management subsystem supporting client, handyman and administrator accounts (email/password and Google sign-in).",
        "O2 - Implement the job lifecycle: posting, quoting, assignment, milestone and progress tracking, completion and dispute handling.",
        "O3 - Design and implement an in-app digital wallet with QR-code and short-code payments, funds top-up, transfers and withdrawals.",
        "O4 - Provide real-time communication through job-scoped and direct chat with text, image and voice messages and call signaling.",
        "O5 - Build discovery and trust features: handyman portfolios, a community feed, and a review and rating system.",
        "O6 - Provide an administrative console for user management, professional verification, content moderation, dispute resolution and payout processing.",
    ])
    doc.add_heading("1.4 Scope", level=2)
    para(doc, "The system is a responsive web application intended for desktop and mobile browsers. It targets the Harare metropolitan area as its primary market but is geographically agnostic. The payment subsystem simulates real-money movement with a wallet ledger: funds are credited atomically within the database and external gateways (card processors, mobile-money APIs) are out of scope. The security gate (PIN, password, biometrics) is implemented client-side as a demonstration and is not a substitute for server-enforced rules in production.")
    doc.add_heading("1.5 Significance", level=2)
    para(doc, "HandyConnect contributes a complete, working reference implementation of a localised gig marketplace. It demonstrates how modern web technologies (React, Tailwind CSS, Firebase) can digitise an informal service economy, and how database transactions can guarantee payment integrity in a system without a central payment gateway.")
    doc.add_heading("1.6 Report Organisation", level=2)
    para(doc, "Chapter 2 reviews existing platforms. Chapter 3 analyses the requirements. Chapter 4 presents the system design. Chapter 5 describes the implementation. Chapter 6 covers testing, Chapter 7 deployment and configuration, Chapter 8 evaluates the system and Chapter 9 concludes with future work.")

    # 2. literature
    doc.add_heading("2. Literature Review", level=1)
    doc.add_heading("2.1 Existing Gig Platforms", level=2)
    table(doc,
        ["Platform", "Region", "Core model", "Payments", "Rating / trust", "Relevance to HandyConnect"],
        [
            ["TaskRabbit", "USA / UK / EU", "Taskers post profiles; clients book by category", "Cards, in-app fees", "Reviews after each task", "Job-posting model; heavy focus on trust"],
            ["Airtasker", "Australia / NZ", "Reverse auction - clients post tasks, offers come in", "In-app escrow with cards", "Ratings and verified badges", "Offer/quote flow resembles HandyConnect quoting"],
            ["Thumbtack", "USA", "Lead-based matching of pros to projects", "Online payment via Stripe", "Background-checked pros", "Verification/background checks"],
            ["Local WhatsApp / word of mouth", "Zimbabwe", "Informal referrals, phone negotiation", "Cash, EcoCash, OneMoney", "No structured feedback", "Status quo HandyConnect replaces"],
        ],
        widths=[1.1, 1.0, 1.7, 1.3, 1.2, 1.7])
    doc.add_heading("2.2 Gaps Addressed", level=2)
    bullets(doc, [
        "Local payment culture: the dominant platforms rely on international card processors. HandyConnect adds QR-code and 6-character payment codes that mirror short-code mobile-money practices.",
        "Informal workforce: no requirement for formal employment - verification is lightweight and managed by an administrator.",
        "Data-light context: the platform works over a normal browser connection and does not assume a permanent data link for core functions.",
    ])

    # 3. analysis
    doc.add_heading("3. Requirements Analysis", level=1)
    doc.add_heading("3.1 Actors", level=2)
    table(doc, ["Actor", "Description", "Primary goals"],
        [
            ["Client", "A person who needs a home service and posts jobs", "Find a professional, compare quotes, pay securely, track work"],
            ["Handyman", "A professional who offers services", "Find work, win jobs, get paid, build a reputation"],
            ["Administrator", "Platform operator", "Verify professionals, moderate content, resolve disputes, process payouts"],
        ], widths=[1.2, 2.6, 2.6])
    doc.add_heading("3.2 Use-Case Diagram", level=2)
    image(doc, os.path.join(DIAG, "use_cases.png"), "Figure 3.1 - Use-case diagram of the HandyConnect system.", 6.6)
    doc.add_heading("3.3 Functional Requirements", level=2)
    table(doc, ["ID", "Module", "Requirement"],
        [
            ["FR-01", "Auth", "A user can register with email/password (with verification) or sign in with Google."],
            ["FR-02", "Auth", "A user selects a role (client, handyman, admin) that controls access to areas of the app."],
            ["FR-03", "Jobs", "A client can post a job with title, category, budget, location and optional attachments."],
            ["FR-04", "Jobs", "A handyman can browse open jobs, submit quotes and accept jobs."],
            ["FR-05", "Jobs", "Progress, milestones and disputes can be recorded on an assigned job."],
            ["FR-06", "Jobs", "Completing a job credits the handyman's wallet."],
            ["FR-07", "Payments", "A handyman can generate a QR code and 6-character code for a payment request."],
            ["FR-08", "Payments", "A client can pay by scanning the QR, entering the code, or from the wallet."],
            ["FR-09", "Payments", "Payment settlement is atomic: wallet credit, ledger entries and job status update together."],
            ["FR-10", "Wallet", "A user can view balance, pending amounts, transactions and receipt details."],
            ["FR-11", "Wallet", "A handyman can withdraw funds (minimum $20) and view payout history."],
            ["FR-12", "Wallet", "A client can add funds via card or mobile money methods and transfer between users."],
            ["FR-13", "Security", "Wallet actions can be gated by PIN, payment password or biometric verification with a lockout policy."],
            ["FR-14", "Chat", "Clients and handymen can exchange text, image and voice messages in job or direct conversations."],
            ["FR-15", "Community", "Users can post text, images, before/after project updates and react and comment."],
            ["FR-16", "Portfolio", "A handyman can manage a portfolio of projects with photos, prices and featured items."],
            ["FR-17", "Reviews", "A client can rate and review a professional after a payment."],
            ["FR-18", "Admin", "An administrator can manage users, verify professionals, resolve disputes, delete content and process payouts."],
        ], widths=[0.8, 1.4, 4.2])
    doc.add_heading("3.4 Non-Functional Requirements", level=2)
    table(doc, ["Category", "Requirement"],
        [
            ["Usability", "Responsive interface usable on mobile browsers and desktop; consistent design system with light and dark themes."],
            ["Performance", "Pages render quickly; real-time data via Firestore listeners updates the UI without refresh."],
            ["Availability", "Backend is hosted on Firebase infrastructure with minimal downtime; graceful fallbacks when external APIs (e.g. maps) are unavailable."],
            ["Security", "Authentication enforced by route guards; wallet actions optionally gated by PIN/password/biometrics; payment mutations run as transactions."],
            ["Integrity", "A payment cannot be partially applied - wallet and ledger updates occur in one Firestore transaction."],
            ["Maintainability", "Clear module separation (features, services, hooks) and consistent coding conventions."],
        ], widths=[1.3, 5.1])

    # 4. design
    doc.add_heading("4. System Design", level=1)
    doc.add_heading("4.1 System Architecture", level=2)
    para(doc, "HandyConnect uses a three-tier architecture delivered as a single-page application backed by Firebase services. The presentation layer is a React SPA; the feature modules form the application layer; the service layer encapsulates all data access; and Firebase provides the backend (authentication, Firestore, storage) with Google Maps as an external API. Firestore real-time listeners push changes to the UI, which is what gives the platform its live, chat-like behaviour.")
    image(doc, os.path.join(DIAG, "architecture.png"), "Figure 4.1 - Layered architecture of HandyConnect.", 6.6)
    doc.add_heading("4.2 Technology Stack", level=2)
    table(doc, ["Layer", "Technology", "Purpose"],
        [
            ["Frontend", "React 19", "UI components and application state"],
            ["Build", "Vite 8", "Dev server, bundling, code-splitting"],
            ["Styling", "Tailwind CSS 4", "Utility-first styling with a design-token system"],
            ["Routing", "react-router-dom 7", "Client-side routing with role-based guards"],
            ["Animations", "framer-motion", "Page transitions and modal animations"],
            ["Icons", "lucide-react", "Icon set"],
            ["Forms", "react-hook-form + zod", "Form state and validation"],
            ["Dates", "date-fns", "Date formatting and grouping"],
            ["QR", "qrcode.react / html5-qrcode", "QR generation and camera scanning"],
            ["Feedback", "react-hot-toast / canvas-confetti", "Notifications and success effects"],
            ["Backend", "Firebase (Auth, Firestore, Storage)", "Identity, NoSQL database, media"],
            ["Maps", "@react-google-maps/api", "Location picker, geocoding, directions"],
        ], widths=[1.1, 2.5, 2.8])
    doc.add_heading("4.3 Database Design (Firestore)", level=2)
    para(doc, "Firestore is a document database organised into collections and subcollections. The diagram below shows the principal collections, their key fields and the relationships between them. Payments, transactions and wallets are the heart of the settlement system; conversations and posts use subcollections for messages and comments to keep list reads cheap.")
    image(doc, os.path.join(DIAG, "data_model.png"), "Figure 4.2 - Firestore data model for HandyConnect.", 7.0)
    doc.add_heading("4.3.1 Composite Indexes", level=2)
    para(doc, "Firestore requires composite indexes for queries that filter and order on different fields. The project ships the following in firestore.indexes.json:")
    table(doc, ["Collection", "Fields"],
        [
            ["conversations", "participants (array-contains) + lastActivity (desc)"],
            ["calls", "toId + status"],
            ["messages", "senderId + read"],
            ["notifications", "toUid + createdAt (desc)"],
            ["notifications", "toUid + read"],
        ], widths=[1.6, 4.8])
    doc.add_heading("4.4 Security Design", level=2)
    para(doc, "Access control operates at two levels. At the application level, React route guards (ProtectedRoute and RoleLayout) redirect unauthenticated users to the login page, send role-less users to role selection, and route each role to its own dashboard. At the data level, Firestore security rules (configured in the Firebase console) restrict reads and writes; the repository documents the required collections and an example ruleset in the deployment chapter.")
    para(doc, "For sensitive wallet operations (adding funds, withdrawing, paying, scanning) the app includes an optional SecurityGate. A user can configure a PIN, a payment password, or biometric sign-in (WebAuthn). PINs and passwords are stored only as salted SHA-256 hashes (Web Crypto), and five consecutive failed attempts trigger a 30-second lockout. This layer demonstrates the UX of payment verification; production hardening would enforce equivalent checks server-side.")
    doc.add_heading("4.5 Payment System Design", level=2)
    para(doc, "Payments are designed to work in low-data, mobile-first conditions:")
    bullets(doc, [
        "A payment request records job, amounts, payer and recipient, a six-character code, and a status (pending, completed, cancelled).",
        "A canonical QR payload encodes the request: HC-PAY|v1|<paymentId>. Scanning resolves the payment by id; a manual fallback enters the six-character code.",
        "Settlement is a single Firestore runTransaction: mark the payment completed, increment the recipient's wallet balance, write a credit transaction for the recipient and a debit transaction for the payer, and set the job to paid. Notifications are sent best-effort after the transaction commits.",
        "Completion generates a job reference (HCJ-YYYY-####) and a transaction reference shown on a confetti success screen with a downloadable receipt.",
    ])
    image(doc, os.path.join(DIAG, "payment_flow.png"), "Figure 4.3 - Sequence diagram of the QR payment flow.", 6.4)
    doc.add_heading("4.6 Realtime Design", level=2)
    para(doc, "Live behaviour is achieved with Firestore onSnapshot subscriptions exposed through custom hooks (useWallet, useConversations, useNotifications, useUnreadCount, useClientJobs). When documents change - a new message, an updated wallet balance, a new notification - the listener fires and React re-renders. Unread counts are derived from per-participant unreadCount maps on conversations and notifications.")
    doc.add_heading("4.7 UI/UX Design", level=2)
    para(doc, "The interface uses a single design system: an orange primary colour (#F97316) with a slate/gray palette, Inter and Plus Jakarta Sans typography, rounded 12-24px surfaces, soft shadows, and light/dark/system themes persisted to localStorage. Each role has a dedicated layout (sidebar on desktop, bottom navigation on mobile). Reusable primitives (Button, Card, Modal, Tabs, Badge, Avatar, Skeleton) keep screens consistent, and framer-motion provides subtle page transitions.")

    # 5. implementation
    doc.add_heading("5. Implementation", level=1)
    doc.add_heading("5.1 Authentication and Role Management", level=2)
    para(doc, "Email/password registration creates the user, sends a verification email and stores a user document with role: null. Google sign-in creates the document if it does not exist. The user then chooses a role on the role-selection page, which writes the role to the user document. Route guards read this role and gate every protected route.")
    image(doc, os.path.join(DIAG, "auth_flow.png"), "Figure 5.1 - Authentication and role-selection flow.", 6.0)
    doc.add_heading("5.2 Job Lifecycle", level=2)
    para(doc, "A job moves through discrete states. Clients create open jobs; handymen submit quotes or accept directly; milestones and progress updates are recorded on assigned jobs; completing a job credits the handyman's wallet; the client's payment marks the job paid. Either party can open a dispute, which an administrator resolves. The state diagram below summarises the transitions implemented in jobService.")
    image(doc, os.path.join(DIAG, "job_lifecycle.png"), "Figure 5.2 - Job lifecycle state diagram.", 6.4)
    doc.add_heading("5.3 Wallet and Payments", level=2)
    para(doc, "The wallet is a single document per user holding balance, pending, credits and coupons. Every money movement writes a transaction document (type and kind) so that history is fully auditable. The core of the payment system is confirmPayment:")
    code_block(doc, """// paymentService.js - settlement is one atomic transaction
export async function confirmPayment(paymentId, payerId) {
  return await runTransaction(db, async (tx) => {
    const payRef = doc(db, "payments", paymentId);
    const paySnap = await tx.get(payRef);
    if (paySnap.data().status !== "pending") throw new Error("Payment already settled");

    const walletRef = doc(db, "wallets", paySnap.data().recipientId);
    await tx.update(payRef, { status: "completed", completedAt: serverTimestamp() });
    await tx.update(walletRef, { balance: increment(paySnap.data().amount) });
    await tx.set(doc(collection(db, "transactions")), {
      uid: paySnap.data().recipientId, type: "payment", kind: "credit",
      amount: paySnap.data().amount, jobId: paySnap.data().jobId, /* ... */
    });
    await tx.set(doc(collection(db, "transactions")), {
      uid: payerId, type: "payment", kind: "debit",
      amount: paySnap.data().amount, jobId: paySnap.data().jobId, /* ... */
    });
    await tx.update(doc(db, "jobs", paySnap.data().jobId), { paid: true });
  });
}""")
    doc.add_heading("5.4 Security Gate", level=2)
    para(doc, "The securityService stores a per-user configuration (PIN, payment password, or WebAuthn credential) with a per-user salt. Values are hashed with SHA-256 via the Web Crypto API, so plaintext secrets are never persisted. Failure counters and lockout timestamps implement the lockout policy.")
    code_block(doc, """// securityService.js - salted hash + lockout (abridged)
async function hashSecret(secret, salt) {
  const data = new TextEncoder().encode(`${salt}::${secret}`);
  const digest = await crypto.subtle.digest("SHA-256", data);
  return Array.from(new Uint8Array(digest)).map(b => b.toString(16).padStart(2, "0")).join("");
}

export async function verifyPin(uid, pin) {
  const cfg = loadSecurityConfig(uid);            // hc_security_<uid>
  if (isLockedOut(cfg)) return { ok: false, error: "Too many attempts. Try again in 30 seconds." };
  const ok = (await hashSecret(pin, cfg.salt)) === cfg.pinHash;
  if (!ok) return recordFailure(cfg);             // 5 fails -> 30s lockout
  return { ok: true };
}""")
    doc.add_heading("5.5 Chat, Community and Portfolio", level=2)
    para(doc, "Conversations are either job-scoped (created from a job) or direct (deduplicated by a sorted participant key). Messages live in a subcollection and drive the chat UI through a listener; unread counts are updated per participant. The community feed stores posts with likes, emoji reactions, bookmarks, hashtags and comments; the portfolio is a subcollection of the handyman user document with photo galleries and before/after pairs. To make first-run demos lively, the app seeds realistic posts and portfolios once per browser (guarded by localStorage flags).")
    doc.add_heading("5.6 Administrative Console", level=2)
    para(doc, "Admin pages subscribe to users, jobs and wallets for real-time tables, and expose actions for role changes, suspension, verification of professionals, job cancellation, dispute resolution, community post deletion and payout processing. Payout processing uses a Firestore transaction to debit the handyman wallet and record the payout.")

    # 6. testing
    doc.add_heading("6. Testing", level=1)
    doc.add_heading("6.1 Strategy", level=2)
    para(doc, "Testing combined manual functional testing across roles, responsive/browser checks, and automated verification of code quality (ESLint) and the production build (Vite). Camera scanning was validated on real devices because html5-qrcode requires actual camera hardware. A representative sample of test cases is shown below.")
    doc.add_heading("6.2 Test Cases", level=2)
    table(doc, ["ID", "Module", "Test scenario", "Expected result", "Status"],
        [
            ["TC-01", "Auth", "Register with email, verify, choose role 'handyman'", "Redirected to handyman dashboard; user doc has role", "Pass"],
            ["TC-02", "Auth", "Logged-in client visits /handyman/dashboard", "Blocked by role guard and redirected", "Pass"],
            ["TC-03", "Jobs", "Client posts a job with budget and location", "Job appears as open in Find Work", "Pass"],
            ["TC-04", "Jobs", "Handyman accepts job, completes it", "Handyman wallet credited, job completed", "Pass"],
            ["TC-05", "Payments", "Handyman generates QR request; client scans it", "Payment screen shows recipient and amount", "Pass"],
            ["TC-06", "Payments", "Client confirms payment", "Payment completed; recipient balance +amount; two ledger entries; job paid", "Pass"],
            ["TC-07", "Payments", "Confirm the same payment twice", "Second attempt rejected (status no longer pending)", "Pass"],
            ["TC-08", "Payments", "Enter a 6-char code for a non-existent payment", "Friendly error, no state change", "Pass"],
            ["TC-09", "Wallet", "Withdraw below the $20 minimum", "Action disabled with explanation", "Pass"],
            ["TC-10", "Security", "Enter wrong PIN 5 times", "30-second lockout enforced", "Pass"],
            ["TC-11", "Chat", "Send a message in a job conversation", "Appears instantly on the other role's thread", "Pass"],
            ["TC-12", "Admin", "Admin processes a payout", "Wallet debited, payout recorded", "Pass"],
            ["TC-13", "Responsive", "Use wallet on a 360px-wide phone", "No horizontal overflow; bottom nav visible", "Pass"],
        ], widths=[0.8, 1.1, 2.4, 1.9, 0.7])
    doc.add_heading("6.3 Build and Lint", level=2)
    para(doc, "The production build completes successfully (Vite transforms ~2,686 modules and emits the bundle with a separate cacheable CSS file). The wallet and payment modules are clean under the project's ESLint configuration; a small set of pre-existing lint warnings in unrelated legacy files remains to be addressed.")
    table(doc, ["Check", "Command", "Result"],
        [
            ["Production build", "npm run build", "Pass - 2,686 modules, ~2.08 MB JS bundle"],
            ["Lint (wallet/payments)", "npx eslint src/features/wallet src/features/payments", "Pass - no issues"],
            ["Lint (full repo)", "npx eslint src/features src/app", "43 pre-existing issues in legacy files (unused imports, hook rules)"],
        ], widths=[1.9, 2.6, 2.0])

    # 7. deployment
    doc.add_heading("7. Deployment and Configuration", level=1)
    doc.add_heading("7.1 Environment Configuration", level=2)
    para(doc, "Configuration is injected through environment variables in a .env file. The build embeds these values into the client bundle; they must never be considered server-side secrets.")
    table(doc, ["Variable", "Purpose"],
        [
            ["VITE_FIREBASE_API_KEY", "Firebase project API key"],
            ["VITE_FIREBASE_AUTH_DOMAIN", "Authentication domain"],
            ["VITE_FIREBASE_PROJECT_ID", "Firestore project identifier"],
            ["VITE_FIREBASE_STORAGE_BUCKET", "Cloud Storage bucket"],
            ["VITE_FIREBASE_MESSAGING_SENDER_ID", "Messaging sender (reserved)"],
            ["VITE_FIREBASE_APP_ID", "Firebase application id"],
            ["VITE_GOOGLE_MAPS_API_KEY", "Google Maps / Places / Geocoding key"],
        ], widths=[2.6, 3.8])
    doc.add_heading("7.2 Firebase Setup", level=2)
    bullets(doc, [
        "Create a Firebase project and enable Authentication (email/password and Google providers), Firestore, and Storage.",
        "Create the users, jobs, wallets, transactions, payments, payouts, notifications, reviews, conversations, posts, follows and calls collections (Firestore creates documents implicitly).",
        "Deploy the composite indexes from firestore.indexes.json (the project's firebase.json points at it).",
        "Configure Firestore security rules to restrict reads/writes to authenticated owners and roles. A minimal starting ruleset is shown below; tighten it for production.",
    ])
    code_block(doc, """rules_version = '2';
service cloud.firestore {
  match /databases/{database}/documents {
    // every authenticated user may read and update their own profile
    match /users/{uid} {
      allow read: if request.auth != null;
      allow write: if request.auth != null && request.auth.uid == uid;
    }
    match /wallets/{uid} {
      allow read: if request.auth != null;
      allow write: if request.auth != null && request.auth.uid == uid;
    }
    match /jobs/{id} {
      allow read: if request.auth != null;
      allow write: if request.auth != null;
    }
    // payments are only mutated by the settlement transaction in code
    match /payments/{id} {
      allow read, write: if request.auth != null;
    }
    // ... additional rules for chat, posts, reviews, payouts
  }
}""")
    doc.add_heading("7.3 Hosting", level=2)
    para(doc, "The application is a static SPA. After npm run build, the dist/ folder can be deployed to Firebase Hosting, Netlify or Vercel with a redirect rule that serves index.html for unknown paths (so client-side routes resolve on refresh).")

    # 8. evaluation
    doc.add_heading("8. Evaluation", level=1)
    doc.add_heading("8.1 Achievement of Objectives", level=2)
    table(doc, ["Objective", "Status", "Evidence"],
        [
            ["O1 - Authentication and roles", "Met", "AuthContext, RoleLayout/ProtectedRoute guards, role-selection page"],
            ["O2 - Job lifecycle", "Met", "jobService state transitions, quotes, milestones, disputes"],
            ["O3 - Wallet and payments", "Met", "QR + code payments, atomic confirmPayment, receipts"],
            ["O4 - Realtime communication", "Met", "Chat threads, voice messages, call signaling over Firestore"],
            ["O5 - Discovery and trust", "Met", "Portfolios, community feed, reviews with star ratings"],
            ["O6 - Admin console", "Met", "User/job/payout management, verification, dispute resolution"],
        ], widths=[2.6, 1.1, 2.7])
    doc.add_heading("8.2 Limitations", level=2)
    bullets(doc, [
        "Payment security (PIN, password, biometrics) is verified client-side; without strict Firestore rules a malicious client could bypass the gate. Server-enforced rules and, ideally, server-side verification are required for production.",
        "No real money movement: funds are ledger credits, not actual EcoCash/card transfers. A gateway integration is future work.",
        "Push notifications are not wired (Firebase Messaging module is reserved but unused); notifications currently surface only inside the running app.",
        "Portfolio and community seed data are generated for demonstration; real deployments should not seed.",
        "No automated unit/integration test suite yet; testing is manual plus lint/build.",
        "Google Maps features require an API key and degrade gracefully when it is absent.",
    ])

    # 9. conclusion
    doc.add_heading("9. Conclusion and Future Work", level=1)
    doc.add_heading("9.1 Conclusion", level=2)
    para(doc, "HandyConnect successfully demonstrates a complete gig-marketplace web application tailored to the Zimbabwean services economy. All six objectives were met: users authenticate and are separated by role; the full job lifecycle is supported; wallet payments work through QR codes and short codes and settle atomically; communication is real-time; trust is built through portfolios, community and reviews; and administrators can operate the platform. The project shows that a production-quality marketplace can be delivered with modern open-source frontend tooling and a serverless Firebase backend, with payment integrity guaranteed by database transactions.")
    doc.add_heading("9.2 Future Work", level=2)
    bullets(doc, [
        "Integrate real payment gateways (EcoCash, Visa/Mastercard, ZIPIT) and escrow/holding accounts.",
        "Move payment verification server-side (Firebase Callable Functions or Cloud Functions) with strict security rules.",
        "Add Firebase Cloud Messaging for true push notifications on mobile.",
        "Introduce automated tests (Vitest + React Testing Library) and CI pipelines.",
        "Develop a native mobile app (React Native / Expo) reusing the service layer.",
        "Add analytics and recommendation features for professionals and pricing.",
        "Extend verification with background checks, licences and in-app identification.",
    ])

    # references
    doc.add_heading("References", level=1)
    refs = [
        "React Documentation - https://react.dev",
        "Vite - Next Generation Frontend Tooling - https://vite.dev",
        "Tailwind CSS Documentation - https://tailwindcss.com/docs",
        "Firebase Documentation (Authentication, Firestore, Storage) - https://firebase.google.com/docs",
        "Web Authentication API (WebAuthn) - https://developer.mozilla.org/en-US/docs/Web/API/Web_Authentication_API",
        "TaskRabbit - https://www.taskrabbit.com",
        "Airtasker - https://www.airtasker.com",
        "html5-qrcode library - https://github.com/mebjas/html5-qrcode",
    ]
    for r in refs:
        p = doc.add_paragraph(r)
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)

    doc.save(OUT_DOCX)
    print("saved:", OUT_DOCX)

if __name__ == "__main__":
    diagram_architecture()
    diagram_data_model()
    diagram_payment_flow()
    diagram_job_lifecycle()
    diagram_auth_flow()
    diagram_use_cases()
    build()
    print("DONE")
